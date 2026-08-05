#!/usr/bin/env python3
"""
Project Document Generator
===========================
Scans project directories, pipeline YAML files, detects CI/CD platform,
analyzes project structure using AI agents, and generates a detailed
Word document with professional formatting.

Supports: GitHub Actions, Azure DevOps, Jenkins (Jenkinsfile + YAML)
Runs on: Linux, Windows, macOS
"""

# Add these imports at the top of the existing script
import shutil
import socket
import os
import sys
import re
import json
import yaml
import platform
import subprocess
import argparse
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Configuration & Constants
# ---------------------------------------------------------------------------

PIPELINE_PATTERNS = [
    # GitHub Actions
    ".github/workflows/*.yml",
    ".github/workflows/*.yaml",
    # Azure DevOps
    "azure-pipelines.yml",
    "azure-pipelines.yaml",
    ".azure-pipelines/*.yml",
    ".azure-pipelines/*.yaml",
    "pipelines/*.yml",
    "pipelines/*.yaml",
    # Jenkins
    "Jenkinsfile",
    "Jenkinsfile.*",
    "jenkins/*.groovy",
    "jenkins/*.yml",
    "jenkins/*.yaml",
    # GitLab CI
    ".gitlab-ci.yml",
    ".gitlab-ci.yaml",
    # Bitbucket
    "bitbucket-pipelines.yml",
    # CircleCI
    ".circleci/config.yml",
    # Travis CI
    ".travis.yml",
    # Generic
    "ci/*.yml",
    "ci/*.yaml",
    ".ci/*.yml",
    ".ci/*.yaml",
]

# File categories for project scanning
FILE_CATEGORIES = {
    "Source Code": [
        ".py", ".js", ".ts", ".java", ".cs", ".go", ".rs", ".rb",
        ".php", ".swift", ".kt", ".scala", ".cpp", ".c", ".h",
        ".hpp", ".m", ".mm", ".r", ".R", ".jl", ".dart", ".lua",
        ".pl", ".pm", ".ex", ".exs", ".clj", ".hs", ".erl",
    ],
    "Web Frontend": [
        ".html", ".htm", ".css", ".scss", ".sass", ".less",
        ".jsx", ".tsx", ".vue", ".svelte", ".astro",
    ],
    "Configuration": [
        ".yml", ".yaml", ".json", ".toml", ".ini", ".cfg",
        ".conf", ".env", ".properties", ".xml", ".hcl",
    ],
    "Infrastructure as Code": [
        ".tf", ".tfvars", ".bicep", ".pulumi",
    ],
    "Containerization": [
        "Dockerfile", "docker-compose.yml", "docker-compose.yaml",
        ".dockerignore",
    ],
    "Documentation": [
        ".md", ".rst", ".txt", ".adoc", ".tex",
    ],
    "Scripts": [
        ".sh", ".bash", ".zsh", ".ps1", ".psm1", ".bat", ".cmd",
    ],
    "Data & Database": [
        ".sql", ".csv", ".tsv", ".parquet", ".avro",
    ],
    "Testing": [
        "_test.go", "_test.py", ".test.js", ".test.ts",
        ".spec.js", ".spec.ts", "_spec.rb",
    ],
    "Build Files": [
        "Makefile", "CMakeLists.txt", "build.gradle",
        "pom.xml", "package.json", "Cargo.toml",
        "go.mod", "requirements.txt", "Pipfile",
        "setup.py", "setup.cfg", "pyproject.toml",
        "Gemfile", "composer.json",
    ],
}

# Table color palette (no silver/black)
TABLE_COLORS = {
    "header": "2E86AB",      # Steel Blue
    "header_text": "FFFFFF",  # White
    "row_even": "F0F7FA",    # Light Blue
    "row_odd": "FFFFFF",     # White
    "accent_1": "A23B72",    # Berry
    "accent_2": "F18F01",    # Orange
    "accent_3": "2E86AB",    # Steel Blue
    "accent_4": "1B4332",    # Forest Green
    "accent_5": "7209B7",    # Purple
    "border": "2E86AB",      # Steel Blue
}

IGNORE_DIRS = {
    ".git", ".svn", ".hg", "__pycache__", "node_modules",
    ".tox", ".mypy_cache", ".pytest_cache", "venv", ".venv",
    "env", ".env", "dist", "build", ".eggs", "*.egg-info",
    ".idea", ".vscode", ".vs", "bin", "obj", "target",
    ".terraform", ".next", ".nuxt", "coverage",
}

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# AI Agent for Project Analysis
# ---------------------------------------------------------------------------

class AIProjectAgent:
    """
    AI Agent that analyzes project structure, infers project name,
    purpose, tech stack, and generates intelligent summaries.
    Uses heuristic analysis (no external API required) with optional
    OpenAI integration.
    """

    def __init__(self, project_path: str, use_openai: bool = False,
                 openai_api_key: str = None):
        self.project_path = Path(project_path)
        self.use_openai = use_openai
        self.openai_api_key = openai_api_key
        self.analysis_cache = {}

    def detect_project_name(self) -> str:
        """Detect project name using multiple heuristics."""
        candidates = []

        # 1. Check package.json
        pkg_json = self.project_path / "package.json"
        if pkg_json.exists():
            try:
                with open(pkg_json, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if "name" in data and data["name"]:
                        candidates.append(("package.json", data["name"], 95))
            except Exception:
                pass

        # 2. Check setup.py / setup.cfg / pyproject.toml
        for setup_file in ["setup.py", "setup.cfg", "pyproject.toml"]:
            fpath = self.project_path / setup_file
            if fpath.exists():
                try:
                    content = fpath.read_text(encoding="utf-8")
                    if setup_file == "pyproject.toml":
                        match = re.search(
                            r'name\s*=\s*["\']([^"\']+)["\']', content
                        )
                    elif setup_file == "setup.cfg":
                        match = re.search(
                            r'name\s*=\s*(.+)', content
                        )
                    else:
                        match = re.search(
                            r'name\s*=\s*["\']([^"\']+)["\']', content
                        )
                    if match:
                        candidates.append(
                            (setup_file, match.group(1).strip(), 90)
                        )
                except Exception:
                    pass

        # 3. Check Cargo.toml
        cargo = self.project_path / "Cargo.toml"
        if cargo.exists():
            try:
                content = cargo.read_text(encoding="utf-8")
                match = re.search(
                    r'name\s*=\s*["\']([^"\']+)["\']', content
                )
                if match:
                    candidates.append(("Cargo.toml", match.group(1), 90))
            except Exception:
                pass

        # 4. Check pom.xml
        pom = self.project_path / "pom.xml"
        if pom.exists():
            try:
                content = pom.read_text(encoding="utf-8")
                match = re.search(
                    r'<artifactId>([^<]+)</artifactId>', content
                )
                if match:
                    candidates.append(("pom.xml", match.group(1), 90))
            except Exception:
                pass

        # 5. Check build.gradle
        gradle = self.project_path / "build.gradle"
        if gradle.exists():
            try:
                content = gradle.read_text(encoding="utf-8")
                match = re.search(r"rootProject\.name\s*=\s*['\"]([^'\"]+)", content)
                if not match:
                    settings = self.project_path / "settings.gradle"
                    if settings.exists():
                        sc = settings.read_text(encoding="utf-8")
                        match = re.search(
                            r"rootProject\.name\s*=\s*['\"]([^'\"]+)", sc
                        )
                if match:
                    candidates.append(("build.gradle", match.group(1), 88))
            except Exception:
                pass

        # 6. Check go.mod
        gomod = self.project_path / "go.mod"
        if gomod.exists():
            try:
                content = gomod.read_text(encoding="utf-8")
                match = re.search(r'module\s+(\S+)', content)
                if match:
                    mod_name = match.group(1).split("/")[-1]
                    candidates.append(("go.mod", mod_name, 88))
            except Exception:
                pass

        # 7. Check .git remote
        try:
            result = subprocess.run(
                ["git", "remote", "get-url", "origin"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                url = result.stdout.strip()
                # Extract repo name from URL
                repo_name = url.rstrip("/").split("/")[-1]
                repo_name = repo_name.replace(".git", "")
                candidates.append(("git remote", repo_name, 85))
        except Exception:
            pass

        # 8. Check README
        for readme in ["README.md", "README.rst", "README.txt", "README"]:
            rpath = self.project_path / readme
            if rpath.exists():
                try:
                    content = rpath.read_text(encoding="utf-8")
                    lines = content.strip().split("\n")
                    for line in lines[:5]:
                        line = line.strip()
                        # Markdown heading
                        match = re.match(r'^#{1,2}\s+(.+)', line)
                        if match:
                            candidates.append(
                                ("README", match.group(1).strip(), 80)
                            )
                            break
                except Exception:
                    pass

        # 9. Fall back to directory name
        dir_name = self.project_path.resolve().name
        candidates.append(("directory", dir_name, 50))

        # Sort by confidence and return the best
        candidates.sort(key=lambda x: x[2], reverse=True)

        if candidates:
            source, name, confidence = candidates[0]
            logger.info(
                f"Detected project name: '{name}' (source: {source}, "
                f"confidence: {confidence}%)"
            )
            return name

        return "Unknown Project"

    def detect_tech_stack(self, file_tree: Dict) -> Dict[str, List[str]]:
        """Detect technology stack from file analysis."""
        tech = {
            "Languages": set(),
            "Frameworks": set(),
            "Databases": set(),
            "CI/CD": set(),
            "Cloud & Infrastructure": set(),
            "Testing": set(),
            "Package Managers": set(),
            "Containerization": set(),
        }

        all_files = []
        self._flatten_files(file_tree, all_files)
        all_filenames = {Path(f).name for f in all_files}
        all_extensions = {Path(f).suffix.lower() for f in all_files}

        # Language detection
        lang_map = {
            ".py": "Python", ".js": "JavaScript", ".ts": "TypeScript",
            ".java": "Java", ".cs": "C#", ".go": "Go", ".rs": "Rust",
            ".rb": "Ruby", ".php": "PHP", ".swift": "Swift",
            ".kt": "Kotlin", ".scala": "Scala", ".cpp": "C++",
            ".c": "C", ".r": "R", ".R": "R", ".jl": "Julia",
            ".dart": "Dart", ".lua": "Lua", ".pl": "Perl",
            ".ex": "Elixir", ".clj": "Clojure", ".hs": "Haskell",
            ".erl": "Erlang",
        }
        for ext, lang in lang_map.items():
            if ext in all_extensions:
                tech["Languages"].add(lang)

        # Framework detection
        framework_indicators = {
            "requirements.txt": self._check_python_frameworks,
            "package.json": self._check_node_frameworks,
            "Gemfile": lambda: {"Ruby on Rails" if (
                self.project_path / "config" / "routes.rb"
            ).exists() else "Ruby"},
            "pom.xml": lambda: {"Spring Boot / Java"},
            "build.gradle": lambda: {"Gradle / Java"},
            "Cargo.toml": lambda: {"Rust"},
            "go.mod": lambda: {"Go"},
            "pubspec.yaml": lambda: {"Flutter / Dart"},
        }
        for indicator, detector in framework_indicators.items():
            if indicator in all_filenames:
                try:
                    result = detector()
                    if result:
                        tech["Frameworks"].update(result)
                except Exception:
                    pass

        # CI/CD detection
        if any(f.startswith(".github/workflows/") for f in all_files):
            tech["CI/CD"].add("GitHub Actions")
        if any("azure-pipelines" in f.lower() for f in all_files):
            tech["CI/CD"].add("Azure DevOps Pipelines")
        if any("Jenkinsfile" in Path(f).name for f in all_files):
            tech["CI/CD"].add("Jenkins")
        if ".gitlab-ci.yml" in all_filenames:
            tech["CI/CD"].add("GitLab CI")
        if ".travis.yml" in all_filenames:
            tech["CI/CD"].add("Travis CI")
        if "bitbucket-pipelines.yml" in all_filenames:
            tech["CI/CD"].add("Bitbucket Pipelines")
        if ".circleci" in {Path(f).parts[0] for f in all_files if "/" in f}:
            tech["CI/CD"].add("CircleCI")

        # Infrastructure
        if any(f.endswith(".tf") for f in all_files):
            tech["Cloud & Infrastructure"].add("Terraform")
        if any(f.endswith(".bicep") for f in all_files):
            tech["Cloud & Infrastructure"].add("Azure Bicep")
        if any("kubernetes" in f.lower() or "k8s" in f.lower()
               for f in all_files):
            tech["Cloud & Infrastructure"].add("Kubernetes")
        if any("helm" in f.lower() for f in all_files):
            tech["Cloud & Infrastructure"].add("Helm")
        if any("ansible" in f.lower() for f in all_files):
            tech["Cloud & Infrastructure"].add("Ansible")

        # Containerization
        if "Dockerfile" in all_filenames:
            tech["Containerization"].add("Docker")
        if any("docker-compose" in f for f in all_filenames):
            tech["Containerization"].add("Docker Compose")

        # Package managers
        pkg_mgr = {
            "package.json": "npm/yarn",
            "requirements.txt": "pip",
            "Pipfile": "pipenv",
            "poetry.lock": "Poetry",
            "Cargo.lock": "Cargo",
            "go.sum": "Go Modules",
            "Gemfile.lock": "Bundler",
            "composer.lock": "Composer",
            "pnpm-lock.yaml": "pnpm",
            "yarn.lock": "Yarn",
            "package-lock.json": "npm",
        }
        for file, mgr in pkg_mgr.items():
            if file in all_filenames:
                tech["Package Managers"].add(mgr)

        # Testing
        test_indicators = {
            "pytest.ini": "pytest",
            ".pytest.ini": "pytest",
            "conftest.py": "pytest",
            "jest.config.js": "Jest",
            "jest.config.ts": "Jest",
            "karma.conf.js": "Karma",
            "cypress.json": "Cypress",
            "cypress.config.js": "Cypress",
            "cypress.config.ts": "Cypress",
            "playwright.config.ts": "Playwright",
            ".rspec": "RSpec",
            "phpunit.xml": "PHPUnit",
        }
        for file, framework in test_indicators.items():
            if file in all_filenames:
                tech["Testing"].add(framework)

        # Convert sets to sorted lists, remove empty
        return {
            k: sorted(v) for k, v in tech.items() if v
        }

    def _check_python_frameworks(self) -> set:
        """Check Python frameworks from requirements.txt."""
        frameworks = set()
        req_file = self.project_path / "requirements.txt"
        if req_file.exists():
            try:
                content = req_file.read_text(encoding="utf-8").lower()
                fw_map = {
                    "django": "Django",
                    "flask": "Flask",
                    "fastapi": "FastAPI",
                    "tornado": "Tornado",
                    "pyramid": "Pyramid",
                    "celery": "Celery",
                    "airflow": "Apache Airflow",
                    "pandas": "Pandas",
                    "numpy": "NumPy",
                    "tensorflow": "TensorFlow",
                    "pytorch": "PyTorch",
                    "torch": "PyTorch",
                    "scikit-learn": "Scikit-learn",
                    "streamlit": "Streamlit",
                    "gradio": "Gradio",
                    "langchain": "LangChain",
                    "transformers": "Hugging Face Transformers",
                }
                for pkg, name in fw_map.items():
                    if pkg in content:
                        frameworks.add(name)
            except Exception:
                pass
        return frameworks

    def _check_node_frameworks(self) -> set:
        """Check Node.js frameworks from package.json."""
        frameworks = set()
        pkg_file = self.project_path / "package.json"
        if pkg_file.exists():
            try:
                with open(pkg_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                all_deps = {}
                all_deps.update(data.get("dependencies", {}))
                all_deps.update(data.get("devDependencies", {}))
                fw_map = {
                    "react": "React",
                    "next": "Next.js",
                    "vue": "Vue.js",
                    "nuxt": "Nuxt.js",
                    "@angular/core": "Angular",
                    "svelte": "Svelte",
                    "express": "Express.js",
                    "nestjs": "NestJS",
                    "@nestjs/core": "NestJS",
                    "fastify": "Fastify",
                    "gatsby": "Gatsby",
                    "electron": "Electron",
                    "tailwindcss": "Tailwind CSS",
                    "bootstrap": "Bootstrap",
                    "material-ui": "Material UI",
                    "@mui/material": "Material UI",
                    "redux": "Redux",
                    "mobx": "MobX",
                    "prisma": "Prisma",
                    "typeorm": "TypeORM",
                    "sequelize": "Sequelize",
                    "mongoose": "Mongoose",
                }
                for pkg, name in fw_map.items():
                    if pkg in all_deps:
                        frameworks.add(name)
            except Exception:
                pass
        return frameworks

    def _flatten_files(self, file_tree: Dict, result: List,
                       prefix: str = ""):
        """Flatten file tree into list of paths."""
        for key, value in file_tree.items():
            path = f"{prefix}/{key}" if prefix else key
            if isinstance(value, dict):
                self._flatten_files(value, result, path)
            else:
                result.append(path)

    def generate_project_summary(
        self, project_name: str, tech_stack: Dict,
        pipeline_info: List[Dict], file_stats: Dict
    ) -> str:
        """Generate an intelligent project summary."""
        languages = tech_stack.get("Languages", [])
        frameworks = tech_stack.get("Frameworks", [])
        cicd = tech_stack.get("CI/CD", [])

        summary_parts = [
            f"**{project_name}** is a software project"
        ]

        if languages:
            if len(languages) == 1:
                summary_parts.append(
                    f" built primarily with {languages[0]}"
                )
            else:
                summary_parts.append(
                    f" built with {', '.join(languages[:-1])} "
                    f"and {languages[-1]}"
                )

        if frameworks:
            summary_parts.append(
                f", utilizing {', '.join(frameworks)}"
            )

        summary_parts.append(".")

        if cicd:
            summary_parts.append(
                f" The project employs {', '.join(cicd)} for "
                f"continuous integration and deployment."
            )

        if pipeline_info:
            total_stages = sum(
                len(p.get("stages", [])) for p in pipeline_info
            )
            total_jobs = sum(
                len(p.get("jobs", [])) for p in pipeline_info
            )
            if total_stages > 0 or total_jobs > 0:
                summary_parts.append(
                    f" The CI/CD pipeline consists of "
                    f"{len(pipeline_info)} pipeline file(s) with "
                    f"{total_stages} stage(s) and {total_jobs} job(s)."
                )

        total_files = sum(file_stats.get("by_category", {}).values())
        total_dirs = file_stats.get("total_directories", 0)
        summary_parts.append(
            f" The project contains {total_files} files across "
            f"{total_dirs} directories."
        )

        infra = tech_stack.get("Cloud & Infrastructure", [])
        containers = tech_stack.get("Containerization", [])
        if infra or containers:
            items = infra + containers
            summary_parts.append(
                f" Infrastructure is managed using "
                f"{', '.join(items)}."
            )

        return "".join(summary_parts)

    def detect_project_type(self, tech_stack: Dict) -> str:
        """Detect the type/category of the project."""
        languages = set(tech_stack.get("Languages", []))
        frameworks = set(tech_stack.get("Frameworks", []))

        if frameworks & {"React", "Vue.js", "Angular", "Svelte",
                         "Next.js", "Nuxt.js", "Gatsby"}:
            if frameworks & {"Express.js", "NestJS", "Fastify",
                             "Django", "Flask", "FastAPI"}:
                return "Full-Stack Web Application"
            return "Frontend Web Application"

        if frameworks & {"Express.js", "NestJS", "Fastify",
                         "Django", "Flask", "FastAPI",
                         "Spring Boot / Java", "Ruby on Rails"}:
            return "Backend Web Application / API"

        if frameworks & {"TensorFlow", "PyTorch", "Scikit-learn",
                         "LangChain", "Hugging Face Transformers"}:
            return "Machine Learning / AI Project"

        if frameworks & {"Apache Airflow", "Celery"}:
            return "Data Pipeline / ETL"

        if frameworks & {"Electron"}:
            return "Desktop Application"

        if frameworks & {"Flutter / Dart"}:
            return "Mobile Application"

        if tech_stack.get("Cloud & Infrastructure"):
            return "Infrastructure / DevOps Project"

        if languages & {"Python"}:
            return "Python Application"
        if languages & {"JavaScript", "TypeScript"}:
            return "JavaScript/TypeScript Application"
        if languages & {"Java"}:
            return "Java Application"
        if languages & {"Go"}:
            return "Go Application"
        if languages & {"Rust"}:
            return "Rust Application"
        if languages & {"C#"}:
            return "C# / .NET Application"

        return "Software Project"


# ---------------------------------------------------------------------------
# Project Scanner
# ---------------------------------------------------------------------------

class ProjectScanner:
    """Scans the project directory for files, pipelines, and metadata."""

    def __init__(self, project_path: str):
        self.project_path = Path(project_path).resolve()
        if not self.project_path.exists():
            raise FileNotFoundError(
                f"Project path does not exist: {self.project_path}"
            )

    def scan_file_tree(self) -> Dict:
        """Scan entire project directory and build a file tree."""
        logger.info(f"Scanning directory: {self.project_path}")
        tree = {}
        self._build_tree(self.project_path, tree)
        return tree

    def _build_tree(self, current_path: Path, tree: Dict):
        """Recursively build file tree."""
        try:
            items = sorted(current_path.iterdir(),
                           key=lambda x: (not x.is_dir(), x.name.lower()))
        except PermissionError:
            return

        for item in items:
            if item.name in IGNORE_DIRS:
                continue
            if item.name.startswith(".") and item.name not in (
                ".github", ".gitlab-ci.yml", ".circleci",
                ".azure-pipelines", ".travis.yml", ".dockerignore",
                ".env.example", ".gitignore",
            ):
                # Include important dot-files/dirs
                if not item.is_dir():
                    rel = str(item.relative_to(self.project_path))
                    tree[rel] = self._get_file_info(item)
                continue

            if item.is_dir():
                subtree = {}
                self._build_tree(item, subtree)
                if subtree:  # Only include non-empty dirs
                    rel = str(item.relative_to(self.project_path))
                    tree[rel] = subtree
            else:
                rel = str(item.relative_to(self.project_path))
                tree[rel] = self._get_file_info(item)

    def _get_file_info(self, file_path: Path) -> Dict:
        """Get file metadata."""
        try:
            stat = file_path.stat()
            return {
                "size": stat.st_size,
                "modified": datetime.fromtimestamp(
                    stat.st_mtime
                ).strftime("%Y-%m-%d %H:%M"),
                "extension": file_path.suffix.lower(),
                "name": file_path.name,
            }
        except Exception:
            return {
                "size": 0,
                "modified": "Unknown",
                "extension": file_path.suffix.lower(),
                "name": file_path.name,
            }

    def get_file_statistics(self, file_tree: Dict) -> Dict:
        """Calculate file statistics."""
        stats = {
            "total_files": 0,
            "total_directories": 0,
            "total_size": 0,
            "by_extension": defaultdict(int),
            "by_category": defaultdict(int),
            "largest_files": [],
        }
        all_files = []
        self._collect_stats(file_tree, stats, all_files)

        # Categorize files
        for finfo in all_files:
            categorized = False
            for category, extensions in FILE_CATEGORIES.items():
                if (finfo["extension"] in extensions or
                        finfo["name"] in extensions):
                    stats["by_category"][category] += 1
                    categorized = True
                    break
            if not categorized:
                stats["by_category"]["Other"] += 1

        # Sort largest files
        all_files.sort(key=lambda x: x["size"], reverse=True)
        stats["largest_files"] = all_files[:10]

        return stats

    def _collect_stats(self, tree: Dict, stats: Dict,
                       all_files: List):
        """Recursively collect statistics."""
        for key, value in tree.items():
            if isinstance(value, dict):
                if any(isinstance(v, dict) for v in value.values()):
                    # This is a directory
                    stats["total_directories"] += 1
                    self._collect_stats(value, stats, all_files)
                elif "name" in value:
                    # This is a file info dict
                    stats["total_files"] += 1
                    stats["total_size"] += value.get("size", 0)
                    ext = value.get("extension", "")
                    if ext:
                        stats["by_extension"][ext] += 1
                    value["path"] = key
                    all_files.append(value)
                else:
                    # Directory with file info children
                    stats["total_directories"] += 1
                    self._collect_stats(value, stats, all_files)
            else:
                stats["total_files"] += 1

    def find_pipeline_files(self) -> List[Path]:
        """Find all pipeline/CI-CD configuration files."""
        pipeline_files = []
        import glob

        for pattern in PIPELINE_PATTERNS:
            full_pattern = str(self.project_path / pattern)
            matches = glob.glob(full_pattern, recursive=False)
            for match in matches:
                p = Path(match)
                if p.is_file() and p not in pipeline_files:
                    pipeline_files.append(p)

        # Also do recursive search for common patterns
        for root, dirs, files in os.walk(self.project_path):
            # Skip ignored directories
            dirs[:] = [
                d for d in dirs
                if d not in IGNORE_DIRS
            ]
            for fname in files:
                fpath = Path(root) / fname
                if fpath in pipeline_files:
                    continue
                if fname in ("Jenkinsfile",):
                    pipeline_files.append(fpath)
                elif fname.endswith((".yml", ".yaml")):
                    # Check if it looks like a pipeline file
                    try:
                        content = fpath.read_text(
                            encoding="utf-8", errors="ignore"
                        )[:2000]
                        pipeline_indicators = [
                            "trigger:", "stages:", "jobs:",
                            "pipeline:", "steps:", "on:",
                            "workflow_dispatch", "push:",
                            "pull_request:", "pool:",
                            "resources:", "strategy:",
                        ]
                        if any(ind in content
                               for ind in pipeline_indicators):
                            rel = fpath.relative_to(self.project_path)
                            rel_str = str(rel)
                            if any(
                                rel_str.startswith(p.split("*")[0])
                                for p in PIPELINE_PATTERNS
                                if "*" in p
                            ) or any(
                                rel_str == p
                                for p in PIPELINE_PATTERNS
                                if "*" not in p
                            ):
                                continue
                            # Only add if it's clearly a pipeline
                            strong_indicators = [
                                "stages:", "jobs:", "pipeline:",
                                "trigger:", "on:", "pool:",
                            ]
                            if sum(1 for si in strong_indicators
                                   if si in content) >= 2:
                                pipeline_files.append(fpath)
                    except Exception:
                        pass

        logger.info(f"Found {len(pipeline_files)} pipeline file(s)")
        return pipeline_files

    def parse_pipeline_file(self, filepath: Path) -> Dict:
        """Parse a pipeline YAML/Jenkinsfile and extract details."""
        result = {
            "file": str(filepath.relative_to(self.project_path)),
            "type": "Unknown",
            "raw_content": "",
            "stages": [],
            "jobs": [],
            "triggers": [],
            "variables": {},
            "environments": [],
            "services": [],
            "artifacts": [],
            "parameters": [],
            "error": None,
        }

        try:
            content = filepath.read_text(encoding="utf-8")
            result["raw_content"] = content
        except Exception as e:
            result["error"] = str(e)
            return result

        # Detect pipeline type
        if "Jenkinsfile" in filepath.name or filepath.suffix == ".groovy":
            result["type"] = "Jenkins"
            self._parse_jenkinsfile(content, result)
        elif ".github/workflows" in str(filepath):
            result["type"] = "GitHub Actions"
            self._parse_github_actions(content, result)
        elif "azure-pipelines" in filepath.name.lower() or \
                ".azure-pipelines" in str(filepath):
            result["type"] = "Azure DevOps"
            self._parse_azure_pipelines(content, result)
        elif ".gitlab-ci" in filepath.name:
            result["type"] = "GitLab CI"
            self._parse_gitlab_ci(content, result)
        else:
            # Try to detect from content
            self._parse_generic_yaml(content, result)

        return result

    def _parse_github_actions(self, content: str, result: Dict):
        """Parse GitHub Actions workflow."""
        try:
            data = yaml.safe_load(content)
            if not isinstance(data, dict):
                return

            result["type"] = "GitHub Actions"

            # Triggers
            on_config = data.get("on", {})
            if isinstance(on_config, str):
                result["triggers"].append(on_config)
            elif isinstance(on_config, list):
                result["triggers"].extend(on_config)
            elif isinstance(on_config, dict):
                result["triggers"].extend(on_config.keys())

            # Name
            result["workflow_name"] = data.get("name", "Unnamed Workflow")

            # Environment variables
            env = data.get("env", {})
            if isinstance(env, dict):
                result["variables"] = env

            # Jobs
            jobs = data.get("jobs", {})
            if isinstance(jobs, dict):
                for job_name, job_config in jobs.items():
                    job_info = {
                        "name": job_name,
                        "runs_on": "",
                        "steps": [],
                        "needs": [],
                        "condition": "",
                        "environment": "",
                    }
                    if isinstance(job_config, dict):
                        job_info["runs_on"] = str(
                            job_config.get("runs-on", "")
                        )
                        job_info["needs"] = job_config.get("needs", [])
                        if isinstance(job_info["needs"], str):
                            job_info["needs"] = [job_info["needs"]]
                        job_info["condition"] = str(
                            job_config.get("if", "")
                        )
                        job_info["environment"] = str(
                            job_config.get("environment", "")
                        )

                        steps = job_config.get("steps", [])
                        for step in steps:
                            if isinstance(step, dict):
                                step_info = {
                                    "name": step.get(
                                        "name",
                                        step.get("uses",
                                                 step.get("run", "")[:50])
                                    ),
                                    "uses": step.get("uses", ""),
                                    "run": step.get("run", ""),
                                }
                                job_info["steps"].append(step_info)

                        if job_info["environment"]:
                            result["environments"].append(
                                job_info["environment"]
                            )

                    result["jobs"].append(job_info)

            # Extract stages from job dependencies
            if result["jobs"]:
                result["stages"] = [
                    j["name"] for j in result["jobs"]
                ]

        except yaml.YAMLError as e:
            result["error"] = f"YAML parse error: {e}"

    def _parse_azure_pipelines(self, content: str, result: Dict):
        """Parse Azure DevOps pipeline."""
        try:
            data = yaml.safe_load(content)
            if not isinstance(data, dict):
                return

            result["type"] = "Azure DevOps"

            # Trigger
            trigger = data.get("trigger", {})
            if isinstance(trigger, list):
                result["triggers"] = trigger
            elif isinstance(trigger, dict):
                branches = trigger.get("branches", {})
                if isinstance(branches, dict):
                    result["triggers"].extend(
                        branches.get("include", [])
                    )
            elif isinstance(trigger, str):
                result["triggers"].append(trigger)

            # PR trigger
            pr = data.get("pr", {})
            if pr:
                result["triggers"].append("Pull Request")

            # Variables
            variables = data.get("variables", {})
            if isinstance(variables, dict):
                result["variables"] = variables
            elif isinstance(variables, list):
                for var in variables:
                    if isinstance(var, dict):
                        if "name" in var:
                            result["variables"][var["name"]] = var.get(
                                "value", ""
                            )
                        elif "group" in var:
                            result["variables"][
                                f"group:{var['group']}"
                            ] = "(Variable Group)"

            # Pool
            pool = data.get("pool", {})
            if isinstance(pool, dict):
                result["pool"] = pool.get("vmImage", str(pool))
            elif isinstance(pool, str):
                result["pool"] = pool

            # Parameters
            params = data.get("parameters", [])
            if isinstance(params, list):
                for param in params:
                    if isinstance(param, dict):
                        result["parameters"].append({
                            "name": param.get("name", ""),
                            "type": param.get("type", "string"),
                            "default": str(param.get("default", "")),
                        })

            # Stages
            stages = data.get("stages", [])
            if isinstance(stages, list):
                for stage in stages:
                    if isinstance(stage, dict):
                        stage_info = {
                            "name": stage.get(
                                "stage",
                                stage.get("displayName", "Unnamed")
                            ),
                            "displayName": stage.get("displayName", ""),
                            "condition": stage.get("condition", ""),
                            "dependsOn": stage.get("dependsOn", []),
                            "jobs": [],
                        }
                        stage_jobs = stage.get("jobs", [])
                        for job in stage_jobs:
                            if isinstance(job, dict):
                                job_info = {
                                    "name": job.get(
                                        "job",
                                        job.get(
                                            "deployment",
                                            job.get("displayName",
                                                    "Unnamed")
                                        )
                                    ),
                                    "displayName": job.get(
                                        "displayName", ""
                                    ),
                                    "pool": str(job.get("pool", "")),
                                    "steps": [],
                                    "environment": str(
                                        job.get("environment", "")
                                    ),
                                }
                                steps = job.get("steps", [])
                                for step in steps:
                                    if isinstance(step, dict):
                                        step_name = step.get(
                                            "displayName",
                                            step.get(
                                                "task",
                                                step.get("script",
                                                         "")[:50]
                                            )
                                        )
                                        job_info["steps"].append({
                                            "name": step_name,
                                            "task": step.get("task", ""),
                                            "script": step.get(
                                                "script", ""
                                            ),
                                        })
                                stage_info["jobs"].append(job_info)
                                result["jobs"].append(job_info)
                                if job_info["environment"]:
                                    result["environments"].append(
                                        job_info["environment"]
                                    )
                        result["stages"].append(stage_info["name"])

            # If no stages, check for direct jobs
            if not stages:
                jobs = data.get("jobs", [])
                if isinstance(jobs, list):
                    for job in jobs:
                        if isinstance(job, dict):
                            job_info = {
                                "name": job.get(
                                    "job",
                                    job.get("displayName", "Unnamed")
                                ),
                                "displayName": job.get(
                                    "displayName", ""
                                ),
                                "steps": [],
                            }
                            steps = job.get("steps", [])
                            for step in steps:
                                if isinstance(step, dict):
                                    job_info["steps"].append({
                                        "name": step.get(
                                            "displayName",
                                            step.get("task", "")
                                        ),
                                    })
                            result["jobs"].append(job_info)

                # If no jobs either, check for direct steps
                if not jobs:
                    steps = data.get("steps", [])
                    if isinstance(steps, list):
                        job_info = {
                            "name": "Default Job",
                            "steps": [],
                        }
                        for step in steps:
                            if isinstance(step, dict):
                                job_info["steps"].append({
                                    "name": step.get(
                                        "displayName",
                                        step.get("task",
                                                 step.get("script",
                                                          "")[:50])
                                    ),
                                })
                        if job_info["steps"]:
                            result["jobs"].append(job_info)

            # Resources
            resources = data.get("resources", {})
            if isinstance(resources, dict):
                for res_type, res_list in resources.items():
                    if isinstance(res_list, list):
                        for res in res_list:
                            if isinstance(res, dict):
                                result["services"].append({
                                    "type": res_type,
                                    "name": str(
                                        res.get(
                                            res_type.rstrip("s"),
                                            res.get("repository",
                                                    res.get("container",
                                                            ""))
                                        )
                                    ),
                                })

        except yaml.YAMLError as e:
            result["error"] = f"YAML parse error: {e}"

    def _parse_jenkinsfile(self, content: str, result: Dict):
        """Parse Jenkinsfile (Groovy-based)."""
        result["type"] = "Jenkins"

        # Extract stages
        stage_pattern = r"stage\s*\(\s*['\"]([^'\"]+)['\"]\s*\)"
        stages = re.findall(stage_pattern, content)
        result["stages"] = stages

        # Extract agent
        agent_pattern = r"agent\s*\{([^}]+)\}"
        agent_match = re.search(agent_pattern, content)
        if agent_match:
            result["agent"] = agent_match.group(1).strip()

        # Extract environment variables
        env_pattern = r"environment\s*\{([^}]+)\}"
        env_match = re.search(env_pattern, content)
        if env_match:
            env_block = env_match.group(1)
            var_pattern = r"(\w+)\s*=\s*['\"]?([^'\"\n]+)['\"]?"
            for var_match in re.finditer(var_pattern, env_block):
                result["variables"][var_match.group(1)] = \
                    var_match.group(2).strip()

        # Extract triggers
        trigger_pattern = r"triggers\s*\{([^}]+)\}"
        trigger_match = re.search(trigger_pattern, content)
        if trigger_match:
            trigger_block = trigger_match.group(1)
            result["triggers"] = [
                t.strip() for t in trigger_block.strip().split("\n")
                if t.strip()
            ]

        # Extract parameters
        param_pattern = r"parameters\s*\{([^}]+)\}"
        param_match = re.search(param_pattern, content)
        if param_match:
            param_block = param_match.group(1)
            p_pattern = (
                r"(string|boolean|choice|text)\s*\("
                r"[^)]*name:\s*['\"]([^'\"]+)['\"]"
            )
            for p_match in re.finditer(p_pattern, param_block):
                result["parameters"].append({
                    "type": p_match.group(1),
                    "name": p_match.group(2),
                })

        # Create jobs from stages
        for stage in stages:
            result["jobs"].append({
                "name": stage,
                "steps": [],
            })

    def _parse_gitlab_ci(self, content: str, result: Dict):
        """Parse GitLab CI configuration."""
        try:
            data = yaml.safe_load(content)
            if not isinstance(data, dict):
                return

            result["type"] = "GitLab CI"

            # Global settings
            stages = data.get("stages", [])
            result["stages"] = stages if isinstance(stages, list) else []

            variables = data.get("variables", {})
            if isinstance(variables, dict):
                result["variables"] = {
                    k: str(v) for k, v in variables.items()
                }

            # Jobs (everything that's not a reserved keyword)
            reserved = {
                "stages", "variables", "image", "services",
                "before_script", "after_script", "cache",
                "include", "default", "workflow", "pages",
            }
            for key, value in data.items():
                if key.startswith(".") or key in reserved:
                    continue
                if isinstance(value, dict):
                    job_info = {
                        "name": key,
                        "stage": value.get("stage", ""),
                        "image": value.get("image", ""),
                        "script": value.get("script", []),
                        "only": value.get("only", []),
                        "except": value.get("except", []),
                        "environment": str(
                            value.get("environment", "")
                        ),
                    }
                    result["jobs"].append(job_info)
                    if job_info["environment"]:
                        result["environments"].append(
                            job_info["environment"]
                        )

            # Services
            services = data.get("services", [])
            if isinstance(services, list):
                for svc in services:
                    if isinstance(svc, str):
                        result["services"].append({"name": svc})
                    elif isinstance(svc, dict):
                        result["services"].append(svc)

        except yaml.YAMLError as e:
            result["error"] = f"YAML parse error: {e}"

    def _parse_generic_yaml(self, content: str, result: Dict):
        """Parse a generic YAML pipeline file."""
        try:
            data = yaml.safe_load(content)
            if not isinstance(data, dict):
                return

            # Try to detect type
            if "on" in data and "jobs" in data:
                self._parse_github_actions(content, result)
            elif "trigger" in data or "pool" in data:
                self._parse_azure_pipelines(content, result)
            elif "stages" in data and "variables" in data:
                self._parse_gitlab_ci(content, result)
            else:
                result["type"] = "Generic YAML Pipeline"
                # Extract what we can
                for key in ["stages", "jobs", "steps", "tasks"]:
                    if key in data:
                        val = data[key]
                        if isinstance(val, list):
                            result["stages"] = [
                                str(item) if not isinstance(item, dict)
                                else item.get(
                                    "name", item.get("stage", str(item))
                                )
                                for item in val
                            ]

        except yaml.YAMLError:
            pass

    def get_git_info(self) -> Dict:
        """Get git repository information."""
        info = {
            "is_git_repo": False,
            "remote_url": "",
            "current_branch": "",
            "last_commit": "",
            "last_commit_author": "",
            "last_commit_date": "",
            "total_commits": 0,
            "contributors": [],
        }

        try:
            # Check if git repo
            result = subprocess.run(
                ["git", "rev-parse", "--is-inside-work-tree"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=10
            )
            if result.returncode != 0:
                return info

            info["is_git_repo"] = True

            # Remote URL
            result = subprocess.run(
                ["git", "remote", "get-url", "origin"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                info["remote_url"] = result.stdout.strip()

            # Current branch
            result = subprocess.run(
                ["git", "branch", "--show-current"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                info["current_branch"] = result.stdout.strip()

            # Last commit
            result = subprocess.run(
                ["git", "log", "-1",
                 "--format=%H|%an|%ad|%s", "--date=short"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                parts = result.stdout.strip().split("|", 3)
                if len(parts) >= 4:
                    info["last_commit"] = parts[3]
                    info["last_commit_author"] = parts[1]
                    info["last_commit_date"] = parts[2]

            # Total commits
            result = subprocess.run(
                ["git", "rev-list", "--count", "HEAD"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=10
            )
            if result.returncode == 0:
                info["total_commits"] = int(result.stdout.strip())

            # Contributors
            result = subprocess.run(
                ["git", "shortlog", "-sn", "--all"],
                cwd=str(self.project_path),
                capture_output=True, text=True, timeout=15
            )
            if result.returncode == 0:
                for line in result.stdout.strip().split("\n")[:10]:
                    line = line.strip()
                    if line:
                        match = re.match(r'\s*(\d+)\s+(.+)', line)
                        if match:
                            info["contributors"].append({
                                "name": match.group(2).strip(),
                                "commits": int(match.group(1)),
                            })

        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass

        return info


# ---------------------------------------------------------------------------
# Document Generator
# ---------------------------------------------------------------------------

class DocumentGenerator:
    """Generates a professional Word document from project analysis."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        styles = self.doc.styles

        # --- Project Title Style ---
        if "ProjectTitle" not in [s.name for s in styles]:
            title_style = styles.add_style(
                "ProjectTitle", WD_STYLE_TYPE.PARAGRAPH
            )
        else:
            title_style = styles["ProjectTitle"]
        title_font = title_style.font
        title_font.name = "Times New Roman"
        title_font.size = Pt(22)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = \
            WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

        # --- Heading Style (Times New Roman 16) ---
        heading1 = styles["Heading 1"]
        h1_font = heading1.font
        h1_font.name = "Times New Roman"
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)
        heading1.paragraph_format.space_before = Pt(18)
        heading1.paragraph_format.space_after = Pt(8)

        # --- Sub Heading Style (Calibri 14) ---
        heading2 = styles["Heading 2"]
        h2_font = heading2.font
        h2_font.name = "Calibri"
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        heading2.paragraph_format.space_before = Pt(14)
        heading2.paragraph_format.space_after = Pt(6)

        # --- Heading 3 ---
        heading3 = styles["Heading 3"]
        h3_font = heading3.font
        h3_font.name = "Calibri"
        h3_font.size = Pt(13)
        h3_font.bold = True
        h3_font.color.rgb = RGBColor(0, 0, 0)

        # --- Content Style (Calibri 12) ---
        normal = styles["Normal"]
        normal_font = normal.font
        normal_font.name = "Calibri"
        normal_font.size = Pt(12)
        normal_font.color.rgb = RGBColor(0, 0, 0)
        normal.paragraph_format.space_after = Pt(4)
        normal.paragraph_format.line_spacing = 1.15

        # --- List Bullet ---
        if "List Bullet" in [s.name for s in styles]:
            lb = styles["List Bullet"]
            lb.font.name = "Calibri"
            lb.font.size = Pt(12)
            lb.font.color.rgb = RGBColor(0, 0, 0)

        # --- Set default section margins ---
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def _add_title(self, text: str):
        """Add project title."""
        para = self.doc.add_paragraph(text, style="ProjectTitle")
        return para

    def _add_heading(self, text: str, level: int = 1):
        """Add heading."""
        return self.doc.add_heading(text, level=level)

    def _add_paragraph(self, text: str, bold: bool = False,
                       italic: bool = False):
        """Add a content paragraph."""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = bold
        run.italic = italic
        return para

    def _add_bullet_list(self, items: List[str], level: int = 0):
        """Add bullet list items."""
        for item in items:
            para = self.doc.add_paragraph(
                style="List Bullet"
            )
            # Set indentation for nested levels
            if level > 0:
                para.paragraph_format.left_indent = Inches(
                    0.5 * (level + 1)
                )
            run = para.add_run(item)
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_numbered_list(self, items: List[str]):
        """Add numbered list items."""
        for i, item in enumerate(items, 1):
            para = self.doc.add_paragraph()
            run = para.add_run(f"{i}. {item}")
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.paragraph_format.left_indent = Inches(0.5)

    def _add_table(self, headers: List[str], rows: List[List[str]],
                   color_scheme: str = "default"):
        """Add a colorful table."""
        table = self.doc.add_table(
            rows=1 + len(rows), cols=len(headers)
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Color schemes
        schemes = {
            "default": {
                "header_bg": "2E86AB",
                "header_fg": "FFFFFF",
                "even_bg": "E8F4F8",
                "odd_bg": "FFFFFF",
            },
            "green": {
                "header_bg": "1B4332",
                "header_fg": "FFFFFF",
                "even_bg": "D8F3DC",
                "odd_bg": "FFFFFF",
            },
            "purple": {
                "header_bg": "7209B7",
                "header_fg": "FFFFFF",
                "even_bg": "F3E5F5",
                "odd_bg": "FFFFFF",
            },
            "orange": {
                "header_bg": "E85D04",
                "header_fg": "FFFFFF",
                "even_bg": "FFF3E0",
                "odd_bg": "FFFFFF",
            },
            "berry": {
                "header_bg": "A23B72",
                "header_fg": "FFFFFF",
                "even_bg": "FCE4EC",
                "odd_bg": "FFFFFF",
            },
            "teal": {
                "header_bg": "006D77",
                "header_fg": "FFFFFF",
                "even_bg": "E0F2F1",
                "odd_bg": "FFFFFF",
            },
        }
        scheme = schemes.get(color_scheme, schemes["default"])

        # Header row
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(header)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(
                scheme["header_fg"]
            )
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set cell background
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{scheme["header_bg"]}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            bg = scheme["even_bg"] if row_idx % 2 == 0 \
                else scheme["odd_bg"]

            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    run = para.add_run(str(cell_text))
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    # Set row background
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading)

        # Auto-fit columns
        for col_idx in range(len(headers)):
            max_width = len(headers[col_idx])
            for row_data in rows:
                if col_idx < len(row_data):
                    max_width = max(max_width, len(str(row_data[col_idx])))
            # Set column width proportionally
            for row in table.rows:
                if col_idx < len(row.cells):
                    row.cells[col_idx].width = Inches(
                        min(max(max_width * 0.12, 1.0), 3.5)
                    )

        self.doc.add_paragraph()  # Spacer
        return table

    def _add_code_block(self, code: str, max_lines: int = 30):
        """Add a code block with monospace font."""
        lines = code.split("\n")
        if len(lines) > max_lines:
            lines = lines[:max_lines]
            lines.append(f"... ({len(code.split(chr(10))) - max_lines}"
                         f" more lines)")

        for line in lines:
            para = self.doc.add_paragraph()
            run = para.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.left_indent = Inches(0.5)
            # Light background
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F5F5F0"/>'
            )
            para._p.get_or_add_pPr().append(shading)

    def _add_horizontal_line(self):
        """Add a horizontal line separator."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        # Add bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="1" '
            f'w:color="2E86AB"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size to human readable."""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"

    def generate(
        self,
        project_name: str,
        project_type: str,
        project_summary: str,
        tech_stack: Dict,
        pipeline_info: List[Dict],
        file_tree: Dict,
        file_stats: Dict,
        git_info: Dict,
        project_path: str,
        output_path: str,
    ):
        """Generate the complete document."""

        # ============================================================
        # COVER PAGE
        # ============================================================
        # Add some spacing before title
        for _ in range(4):
            self.doc.add_paragraph()

        self._add_title("PROJECT DOCUMENTATION")
        self._add_horizontal_line()

        # Project name subtitle
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(project_name.upper())
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Project type
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(project_type)
        run.font.name = "Calibri"
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Date and system info
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            f"\nGenerated on: {datetime.now().strftime('%B %d, %Y %H:%M')}"
            f"\nPlatform: {platform.system()} {platform.release()}"
            f"\nProject Path: {project_path}"
        )
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

        self.doc.add_page_break()

        # ============================================================
        # TABLE OF CONTENTS (Manual)
        # ============================================================
        self._add_heading("Table of Contents", level=1)
        toc_items = [
            "1. Executive Summary",
            "2. Project Overview",
            "   2.1 Project Information",
            "   2.2 Technology Stack",
            "   2.3 Repository Information",
            "3. Project Structure",
            "   3.1 Directory Overview",
            "   3.2 File Statistics",
            "   3.3 File Categories",
            "   3.4 Largest Files",
            "4. CI/CD Pipeline Analysis",
            "   4.1 Pipeline Overview",
            "   4.2 Detailed Pipeline Configuration",
            "   4.3 Pipeline Stages & Jobs",
            "   4.4 Pipeline Variables & Parameters",
            "5. Environment Configuration",
            "6. Complete File Listing",
            "7. Appendix",
        ]
        for item in toc_items:
            para = self.doc.add_paragraph()
            run = para.add_run(item)
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            if not item.startswith("   "):
                run.bold = True

        self.doc.add_page_break()

        # ============================================================
        # 1. EXECUTIVE SUMMARY
        # ============================================================
        self._add_heading("1. Executive Summary", level=1)
        self._add_paragraph(project_summary)
        self._add_horizontal_line()

        # ============================================================
        # 2. PROJECT OVERVIEW
        # ============================================================
        self._add_heading("2. Project Overview", level=1)

        # 2.1 Project Information
        self._add_heading("2.1 Project Information", level=2)
        info_data = [
            ["Project Name", project_name],
            ["Project Type", project_type],
            ["Project Path", str(project_path)],
            ["Operating System", f"{platform.system()} "
                                 f"{platform.release()}"],
            [
                "Document Generated",
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            ],
            ["Total Files", str(file_stats.get("total_files", 0))],
            [
                "Total Directories",
                str(file_stats.get("total_directories", 0)),
            ],
            [
                "Total Size",
                self._format_file_size(
                    file_stats.get("total_size", 0)
                ),
            ],
        ]
        self._add_table(
            ["Property", "Value"], info_data, "teal"
        )

        # 2.2 Technology Stack
        self._add_heading("2.2 Technology Stack", level=2)
        if tech_stack:
            color_rotation = [
                "default", "green", "purple", "orange",
                "berry", "teal",
            ]
            for idx, (category, items) in enumerate(
                tech_stack.items()
            ):
                self._add_heading(f"  {category}", level=3)
                self._add_bullet_list(items)
        else:
            self._add_paragraph(
                "No specific technology stack detected."
            )

        # 2.3 Repository Information
        self._add_heading("2.3 Repository Information", level=2)
        if git_info.get("is_git_repo"):
            git_data = []
            if git_info["remote_url"]:
                git_data.append(
                    ["Remote URL", git_info["remote_url"]]
                )
            if git_info["current_branch"]:
                git_data.append(
                    ["Current Branch", git_info["current_branch"]]
                )
            if git_info["last_commit"]:
                git_data.append(
                    ["Last Commit", git_info["last_commit"]]
                )
            if git_info["last_commit_author"]:
                git_data.append(
                    ["Last Author", git_info["last_commit_author"]]
                )
            if git_info["last_commit_date"]:
                git_data.append(
                    ["Last Commit Date", git_info["last_commit_date"]]
                )
            git_data.append(
                ["Total Commits", str(git_info.get("total_commits", 0))]
            )

            if git_data:
                self._add_table(
                    ["Property", "Value"], git_data, "green"
                )

            # Contributors table
            if git_info.get("contributors"):
                self._add_heading("Contributors", level=3)
                contrib_rows = [
                    [c["name"], str(c["commits"])]
                    for c in git_info["contributors"]
                ]
                self._add_table(
                    ["Contributor", "Commits"],
                    contrib_rows, "purple"
                )
        else:
            self._add_paragraph(
                "This project is not tracked by Git or Git is not "
                "available."
            )

        self.doc.add_page_break()

        # ============================================================
        # 3. PROJECT STRUCTURE
        # ============================================================
        self._add_heading("3. Project Structure", level=1)

        # 3.1 Directory Overview
        self._add_heading("3.1 Directory Overview", level=2)
        self._add_paragraph(
            "The following is the directory structure of the project:"
        )
        tree_lines = []
        self._render_tree(file_tree, tree_lines, "", True)
        if tree_lines:
            # Show as code block (limited lines)
            tree_text = "\n".join(tree_lines[:100])
            if len(tree_lines) > 100:
                tree_text += (
                    f"\n... ({len(tree_lines) - 100} more entries)"
                )
            self._add_code_block(tree_text, max_lines=100)

        # 3.2 File Statistics
        self._add_heading("3.2 File Statistics", level=2)
        stats_summary = [
            f"Total Files: {file_stats.get('total_files', 0)}",
            f"Total Directories: "
            f"{file_stats.get('total_directories', 0)}",
            f"Total Size: "
            f"{self._format_file_size(file_stats.get('total_size', 0))}",
        ]
        self._add_bullet_list(stats_summary)

        # File extensions table
        ext_stats = file_stats.get("by_extension", {})
        if ext_stats:
            self._add_heading("Files by Extension", level=3)
            ext_sorted = sorted(
                ext_stats.items(), key=lambda x: x[1], reverse=True
            )[:20]
            ext_rows = [
                [ext if ext else "(no extension)", str(count)]
                for ext, count in ext_sorted
            ]
            self._add_table(
                ["Extension", "Count"], ext_rows, "orange"
            )

        # 3.3 File Categories
        self._add_heading("3.3 File Categories", level=2)
        cat_stats = file_stats.get("by_category", {})
        if cat_stats:
            cat_sorted = sorted(
                cat_stats.items(), key=lambda x: x[1], reverse=True
            )
            cat_rows = [
                [cat, str(count)] for cat, count in cat_sorted
            ]
            self._add_table(
                ["Category", "File Count"], cat_rows, "berry"
            )

        # 3.4 Largest Files
        self._add_heading("3.4 Largest Files", level=2)
        largest = file_stats.get("largest_files", [])
        if largest:
            large_rows = [
                [
                    f.get("path", f.get("name", "")),
                    self._format_file_size(f.get("size", 0)),
                    f.get("modified", ""),
                ]
                for f in largest[:10]
            ]
            self._add_table(
                ["File Path", "Size", "Last Modified"],
                large_rows, "teal"
            )

        self.doc.add_page_break()

"""
Project Document Generator - Enhanced Version
=============================================
Added Section 4: Local Development & Deployment Guide
- How to run the project locally (Linux/Windows/macOS)
- How to push to Azure DevOps / GitHub / Jenkins
- Environment setup instructions
- Git workflow guide
- Docker run instructions
- Platform-specific commands
"""
# ---------------------------------------------------------------------------
# Local Development & Deployment Detector
# ---------------------------------------------------------------------------

class LocalDevDeploymentDetector:
    """
    Detects how to run the project locally and how to push/deploy
    to DevOps environments. Analyzes project files to generate
    accurate, platform-specific instructions.
    """

    def __init__(self, project_path: str):
        self.project_path = Path(project_path).resolve()
        self.detected_info = {}

    # ----------------------------------------------------------------
    # LOCAL RUN DETECTION
    # ----------------------------------------------------------------

    def detect_local_run_commands(self) -> Dict[str, Any]:
        """
        Detect how to run the project locally across
        Linux, Windows, and macOS.
        """
        commands = {
            "prerequisites": [],
            "environment_setup": {},
            "install_dependencies": {},
            "run_commands": {},
            "test_commands": {},
            "build_commands": {},
            "docker_commands": {},
            "environment_variables": [],
            "ports": [],
            "notes": [],
        }

        # Detect project runtime
        self._detect_python_commands(commands)
        self._detect_node_commands(commands)
        self._detect_java_commands(commands)
        self._detect_go_commands(commands)
        self._detect_dotnet_commands(commands)
        self._detect_ruby_commands(commands)
        self._detect_php_commands(commands)
        self._detect_rust_commands(commands)
        self._detect_docker_commands(commands)
        self._detect_makefile_commands(commands)
        self._detect_env_variables(commands)
        self._detect_ports(commands)

        return commands

    def _detect_python_commands(self, commands: Dict):
        """Detect Python project run commands."""
        req_file = self.project_path / "requirements.txt"
        pipfile = self.project_path / "Pipfile"
        pyproject = self.project_path / "pyproject.toml"
        setup_py = self.project_path / "setup.py"
        manage_py = self.project_path / "manage.py"

        # Check if it's a Python project
        is_python = any([
            req_file.exists(), pipfile.exists(),
            pyproject.exists(), setup_py.exists(),
            list(self.project_path.glob("*.py")),
        ])

        if not is_python:
            return

        commands["prerequisites"].extend([
            "Python 3.8+ installed",
            "pip package manager",
        ])

        # Detect package manager
        if pipfile.exists():
            pkg_mgr = "pipenv"
            commands["environment_setup"] = {
                "Linux/macOS": [
                    "# Install pipenv",
                    "pip install pipenv",
                    "",
                    "# Create virtual environment",
                    "pipenv install",
                    "",
                    "# Activate virtual environment",
                    "pipenv shell",
                ],
                "Windows (PowerShell)": [
                    "# Install pipenv",
                    "pip install pipenv",
                    "",
                    "# Create virtual environment",
                    "pipenv install",
                    "",
                    "# Activate virtual environment",
                    "pipenv shell",
                ],
            }
            commands["install_dependencies"] = {
                "Linux/macOS": ["pipenv install --dev"],
                "Windows": ["pipenv install --dev"],
            }
        elif pyproject.exists():
            content = pyproject.read_text(encoding="utf-8")
            if "poetry" in content.lower():
                commands["environment_setup"] = {
                    "Linux/macOS": [
                        "# Install Poetry",
                        "curl -sSL https://install.python-poetry.org | python3 -",
                        "",
                        "# Install dependencies",
                        "poetry install",
                        "",
                        "# Activate shell",
                        "poetry shell",
                    ],
                    "Windows (PowerShell)": [
                        "# Install Poetry",
                        "(Invoke-WebRequest -Uri https://install.python-poetry.org "
                        "-UseBasicParsing).Content | python -",
                        "",
                        "# Install dependencies",
                        "poetry install",
                        "",
                        "# Activate shell",
                        "poetry shell",
                    ],
                }
                commands["install_dependencies"] = {
                    "Linux/macOS": ["poetry install --with dev"],
                    "Windows": ["poetry install --with dev"],
                }
            else:
                self._add_venv_setup(commands)
        else:
            self._add_venv_setup(commands)

        # Django project
        if manage_py.exists():
            commands["prerequisites"].append("Django framework")
            commands["run_commands"] = {
                "Linux/macOS": [
                    "# Run database migrations",
                    "python manage.py migrate",
                    "",
                    "# Create superuser (optional)",
                    "python manage.py createsuperuser",
                    "",
                    "# Start development server",
                    "python manage.py runserver",
                    "",
                    "# Start on specific port",
                    "python manage.py runserver 0.0.0.0:8000",
                ],
                "Windows (PowerShell)": [
                    "# Run database migrations",
                    "python manage.py migrate",
                    "",
                    "# Create superuser (optional)",
                    "python manage.py createsuperuser",
                    "",
                    "# Start development server",
                    "python manage.py runserver",
                    "",
                    "# Start on specific port",
                    "python manage.py runserver 0.0.0.0:8000",
                ],
            }
            commands["test_commands"] = {
                "Linux/macOS": [
                    "python manage.py test",
                    "# With coverage",
                    "coverage run manage.py test && coverage report",
                ],
                "Windows": [
                    "python manage.py test",
                    "coverage run manage.py test && coverage report",
                ],
            }
            commands["ports"].append(
                {"port": "8000", "service": "Django Dev Server"}
            )
            return

        # Detect main entry point
        entry_files = ["main.py", "app.py", "run.py", "server.py",
                       "wsgi.py", "asgi.py", "cli.py", "__main__.py"]
        found_entry = None
        for ef in entry_files:
            if (self.project_path / ef).exists():
                found_entry = ef
                break

        # FastAPI / Uvicorn
        req_content = ""
        if req_file.exists():
            req_content = req_file.read_text(
                encoding="utf-8"
            ).lower()

        if "fastapi" in req_content or "uvicorn" in req_content:
            app_module = "main:app" if (
                self.project_path / "main.py"
            ).exists() else "app:app"
            commands["run_commands"] = {
                "Linux/macOS": [
                    "# Development server with auto-reload",
                    f"uvicorn {app_module} --reload",
                    "",
                    "# With specific host and port",
                    f"uvicorn {app_module} --host 0.0.0.0 --port 8000 --reload",
                    "",
                    "# Production",
                    f"uvicorn {app_module} --host 0.0.0.0 --port 8000 "
                    f"--workers 4",
                ],
                "Windows (PowerShell)": [
                    "# Development server with auto-reload",
                    f"uvicorn {app_module} --reload",
                    "",
                    "# With specific host and port",
                    f"uvicorn {app_module} --host 0.0.0.0 --port 8000 --reload",
                ],
            }
            commands["ports"].append(
                {"port": "8000", "service": "FastAPI Server"}
            )

        elif "flask" in req_content:
            commands["run_commands"] = {
                "Linux/macOS": [
                    "# Set Flask app",
                    "export FLASK_APP=app.py",
                    "export FLASK_ENV=development",
                    "",
                    "# Run Flask",
                    "flask run",
                    "",
                    "# Run on specific port",
                    "flask run --host=0.0.0.0 --port=5000",
                ],
                "Windows (PowerShell)": [
                    "# Set Flask app",
                    "$env:FLASK_APP = 'app.py'",
                    "$env:FLASK_ENV = 'development'",
                    "",
                    "# Run Flask",
                    "flask run",
                    "",
                    "# Run on specific port",
                    "flask run --host=0.0.0.0 --port=5000",
                ],
                "Windows (CMD)": [
                    "set FLASK_APP=app.py",
                    "set FLASK_ENV=development",
                    "flask run",
                ],
            }
            commands["ports"].append(
                {"port": "5000", "service": "Flask Dev Server"}
            )

        elif found_entry:
            commands["run_commands"] = {
                "Linux/macOS": [
                    f"python {found_entry}",
                    "",
                    "# With arguments (if applicable)",
                    f"python {found_entry} --help",
                ],
                "Windows (PowerShell)": [
                    f"python {found_entry}",
                    f"python {found_entry} --help",
                ],
            }

        # Test commands
        if "pytest" in req_content or (
            self.project_path / "pytest.ini"
        ).exists() or (
            self.project_path / "conftest.py"
        ).exists():
            commands["test_commands"] = {
                "Linux/macOS": [
                    "# Run all tests",
                    "pytest",
                    "",
                    "# Run with verbose output",
                    "pytest -v",
                    "",
                    "# Run with coverage report",
                    "pytest --cov=. --cov-report=html",
                    "",
                    "# Run specific test file",
                    "pytest tests/test_main.py -v",
                ],
                "Windows": [
                    "pytest",
                    "pytest -v",
                    "pytest --cov=. --cov-report=html",
                ],
            }

    def _add_venv_setup(self, commands: Dict):
        """Add standard venv setup commands."""
        req_file = self.project_path / "requirements.txt"
        req_dev = self.project_path / "requirements-dev.txt"

        install_linux = [
            "# Create virtual environment",
            "python3 -m venv venv",
            "",
            "# Activate virtual environment",
            "source venv/bin/activate",
            "",
        ]
        install_win = [
            "# Create virtual environment",
            "python -m venv venv",
            "",
            "# Activate virtual environment (PowerShell)",
            r".\venv\Scripts\Activate.ps1",
            "",
            "# Activate virtual environment (CMD)",
            r"venv\Scripts\activate.bat",
            "",
        ]

        if req_file.exists():
            install_linux.extend([
                "# Install dependencies",
                "pip install -r requirements.txt",
            ])
            install_win.extend([
                "# Install dependencies",
                "pip install -r requirements.txt",
            ])
            if req_dev.exists():
                install_linux.extend([
                    "",
                    "# Install dev dependencies",
                    "pip install -r requirements-dev.txt",
                ])
                install_win.extend([
                    "",
                    "# Install dev dependencies",
                    "pip install -r requirements-dev.txt",
                ])

        commands["environment_setup"] = {
            "Linux/macOS": install_linux,
            "Windows (PowerShell)": install_win,
        }
        commands["install_dependencies"] = {
            "Linux/macOS": ["pip install -r requirements.txt"],
            "Windows": ["pip install -r requirements.txt"],
        }

    def _detect_node_commands(self, commands: Dict):
        """Detect Node.js project run commands."""
        pkg_json = self.project_path / "package.json"
        if not pkg_json.exists():
            return

        try:
            with open(pkg_json, "r", encoding="utf-8") as f:
                pkg = json.load(f)
        except Exception:
            return

        commands["prerequisites"].extend([
            "Node.js 16+ installed",
            "npm or yarn package manager",
        ])

        # Detect package manager
        has_yarn = (self.project_path / "yarn.lock").exists()
        has_pnpm = (self.project_path / "pnpm-lock.yaml").exists()

        if has_pnpm:
            pkg_cmd = "pnpm"
            install_cmd = "pnpm install"
        elif has_yarn:
            pkg_cmd = "yarn"
            install_cmd = "yarn install"
        else:
            pkg_cmd = "npm"
            install_cmd = "npm install"

        commands["environment_setup"] = {
            "Linux/macOS": [
                "# Install nvm (Node Version Manager) - recommended",
                "curl -o- https://raw.githubusercontent.com/nvm-sh/"
                "nvm/v0.39.0/install.sh | bash",
                "",
                "# Install Node.js",
                "nvm install --lts",
                "nvm use --lts",
                "",
                f"# Install dependencies",
                install_cmd,
            ],
            "Windows (PowerShell)": [
                "# Install nvm-windows from:",
                "# https://github.com/coreybutler/nvm-windows",
                "",
                "# Install Node.js",
                "nvm install lts",
                "nvm use lts",
                "",
                f"# Install dependencies",
                install_cmd,
            ],
            "macOS (Homebrew)": [
                "# Install Node.js via Homebrew",
                "brew install node",
                "",
                f"# Install dependencies",
                install_cmd,
            ],
        }

        commands["install_dependencies"] = {
            "Linux/macOS/Windows": [install_cmd],
        }

        # Scripts from package.json
        scripts = pkg.get("scripts", {})
        run_cmds = {}

        start_script = scripts.get("start", "")
        dev_script = scripts.get("dev", "")
        build_script = scripts.get("build", "")
        test_script = scripts.get("test", "")

        run_list = []
        if dev_script:
            run_list.extend([
                f"# Start development server",
                f"{pkg_cmd} run dev",
                "",
            ])
        if start_script:
            run_list.extend([
                f"# Start production server",
                f"{pkg_cmd} start" if pkg_cmd != "npm"
                else "npm start",
                "",
            ])
        if not run_list:
            run_list = [f"{pkg_cmd} start"]

        commands["run_commands"] = {
            "Linux/macOS": run_list,
            "Windows (PowerShell)": run_list,
        }

        if build_script:
            commands["build_commands"] = {
                "Linux/macOS/Windows": [
                    f"# Build for production",
                    f"{pkg_cmd} run build",
                ],
            }

        if test_script:
            commands["test_commands"] = {
                "Linux/macOS/Windows": [
                    f"# Run tests",
                    f"{pkg_cmd} test",
                    "",
                    f"# Run tests with coverage",
                    f"{pkg_cmd} run test:coverage"
                    if "test:coverage" in scripts
                    else f"{pkg_cmd} test -- --coverage",
                    "",
                    f"# Run tests in watch mode",
                    f"{pkg_cmd} run test:watch"
                    if "test:watch" in scripts
                    else f"{pkg_cmd} test -- --watch",
                ],
            }

        # Check for Next.js
        deps = {}
        deps.update(pkg.get("dependencies", {}))
        deps.update(pkg.get("devDependencies", {}))
        if "next" in deps:
            commands["ports"].append(
                {"port": "3000", "service": "Next.js Dev Server"}
            )
        elif "react-scripts" in deps:
            commands["ports"].append(
                {"port": "3000", "service": "React Dev Server"}
            )
        elif "@angular/core" in deps:
            commands["ports"].append(
                {"port": "4200", "service": "Angular Dev Server"}
            )
        elif "vue" in deps:
            commands["ports"].append(
                {"port": "5173", "service": "Vue/Vite Dev Server"}
            )

        commands["notes"].append(
            f"All scripts defined in package.json: "
            f"{', '.join(scripts.keys())}"
        )

    def _detect_java_commands(self, commands: Dict):
        """Detect Java/Maven/Gradle run commands."""
        pom = self.project_path / "pom.xml"
        gradle = self.project_path / "build.gradle"
        gradlew = self.project_path / "gradlew"
        mvnw = self.project_path / "mvnw"

        if not (pom.exists() or gradle.exists()):
            return

        commands["prerequisites"].extend([
            "Java 11+ (or version specified in project) installed",
            "JAVA_HOME environment variable configured",
        ])

        if pom.exists():
            mvn_cmd = "./mvnw" if mvnw.exists() else "mvn"
            mvn_cmd_win = "mvnw.cmd" if mvnw.exists() else "mvn"

            commands["environment_setup"] = {
                "Linux/macOS": [
                    "# Install Java (Ubuntu/Debian)",
                    "sudo apt-get update",
                    "sudo apt-get install -y openjdk-17-jdk",
                    "",
                    "# Set JAVA_HOME",
                    "export JAVA_HOME=/usr/lib/jvm/java-17-openjdk-amd64",
                    "export PATH=$JAVA_HOME/bin:$PATH",
                    "",
                    "# Verify installation",
                    "java -version",
                    "mvn -version",
                ],
                "Windows (PowerShell)": [
                    "# Install Java via Chocolatey",
                    "choco install openjdk17",
                    "",
                    "# Or download from:",
                    "# https://adoptium.net/",
                    "",
                    "# Set JAVA_HOME in System Environment Variables",
                    '$env:JAVA_HOME = "C:\\Program Files\\Java\\jdk-17"',
                ],
                "macOS (Homebrew)": [
                    "brew install openjdk@17",
                    'sudo ln -sfn /opt/homebrew/opt/openjdk@17/libexec/openjdk.jdk'
                    ' /Library/Java/JavaVirtualMachines/openjdk-17.jdk',
                ],
            }
            commands["install_dependencies"] = {
                "Linux/macOS": [
                    f"{mvn_cmd} dependency:resolve",
                    f"{mvn_cmd} dependency:resolve-sources",
                ],
                "Windows": [
                    f"{mvn_cmd_win} dependency:resolve",
                ],
            }
            commands["build_commands"] = {
                "Linux/macOS": [
                    f"# Clean and build",
                    f"{mvn_cmd} clean install",
                    "",
                    f"# Skip tests during build",
                    f"{mvn_cmd} clean install -DskipTests",
                    "",
                    f"# Package as JAR",
                    f"{mvn_cmd} clean package",
                ],
                "Windows": [
                    f"{mvn_cmd_win} clean install",
                    f"{mvn_cmd_win} clean install -DskipTests",
                    f"{mvn_cmd_win} clean package",
                ],
            }
            commands["run_commands"] = {
                "Linux/macOS": [
                    f"# Build and run",
                    f"{mvn_cmd} spring-boot:run",
                    "",
                    "# Or run the JAR directly",
                    "java -jar target/*.jar",
                    "",
                    "# Run with specific profile",
                    "java -jar target/*.jar --spring.profiles.active=dev",
                ],
                "Windows (PowerShell)": [
                    f"{mvn_cmd_win} spring-boot:run",
                    "java -jar target/*.jar",
                ],
            }
            commands["test_commands"] = {
                "Linux/macOS": [
                    f"{mvn_cmd} test",
                    "",
                    "# Run specific test",
                    f"{mvn_cmd} test -Dtest=TestClassName",
                ],
                "Windows": [f"{mvn_cmd_win} test"],
            }
            commands["ports"].append(
                {"port": "8080", "service": "Spring Boot Server"}
            )

        elif gradle.exists():
            gradle_cmd = "./gradlew" if gradlew.exists() else "gradle"
            gradle_cmd_win = (
                "gradlew.bat" if gradlew.exists() else "gradle"
            )

            commands["build_commands"] = {
                "Linux/macOS": [
                    f"{gradle_cmd} build",
                    f"{gradle_cmd} build -x test",
                ],
                "Windows": [
                    f"{gradle_cmd_win} build",
                ],
            }
            commands["run_commands"] = {
                "Linux/macOS": [
                    f"{gradle_cmd} bootRun",
                    "",
                    "# Run built JAR",
                    "java -jar build/libs/*.jar",
                ],
                "Windows": [
                    f"{gradle_cmd_win} bootRun",
                ],
            }
            commands["test_commands"] = {
                "Linux/macOS": [f"{gradle_cmd} test"],
                "Windows": [f"{gradle_cmd_win} test"],
            }

    def _detect_go_commands(self, commands: Dict):
        """Detect Go project run commands."""
        gomod = self.project_path / "go.mod"
        if not gomod.exists():
            return

        commands["prerequisites"].extend([
            "Go 1.19+ installed",
            "GOPATH configured",
        ])

        # Find main.go
        main_files = list(self.project_path.glob("**/main.go"))
        main_path = (
            str(main_files[0].parent.relative_to(self.project_path))
            if main_files else "."
        )

        commands["environment_setup"] = {
            "Linux/macOS": [
                "# Install Go",
                "# Download from https://go.dev/dl/",
                "",
                "# Or via package manager",
                "sudo apt-get install golang-go  # Ubuntu/Debian",
                "brew install go  # macOS",
                "",
                "# Verify",
                "go version",
                "",
                "# Download dependencies",
                "go mod download",
            ],
            "Windows (PowerShell)": [
                "# Download installer from https://go.dev/dl/",
                "# Or via Chocolatey:",
                "choco install golang",
                "",
                "# Download dependencies",
                "go mod download",
            ],
        }
        commands["install_dependencies"] = {
            "Linux/macOS/Windows": [
                "go mod download",
                "go mod tidy",
            ],
        }
        commands["run_commands"] = {
            "Linux/macOS": [
                f"# Run directly",
                f"go run {main_path}/main.go"
                if main_path != "." else "go run .",
                "",
                f"# Build binary",
                "go build -o app .",
                "",
                "# Run binary",
                "./app",
            ],
            "Windows (PowerShell)": [
                "go run .",
                "",
                "# Build",
                "go build -o app.exe .",
                "",
                "# Run",
                ".\\app.exe",
            ],
        }
        commands["test_commands"] = {
            "Linux/macOS/Windows": [
                "go test ./...",
                "",
                "# With verbose output",
                "go test -v ./...",
                "",
                "# With coverage",
                "go test -cover ./...",
                "go test -coverprofile=coverage.out ./...",
                "go tool cover -html=coverage.out",
            ],
        }
        commands["build_commands"] = {
            "Linux": [
                "GOOS=linux GOARCH=amd64 go build -o app .",
            ],
            "Windows cross-compile from Linux": [
                "GOOS=windows GOARCH=amd64 go build -o app.exe .",
            ],
            "macOS": [
                "GOOS=darwin GOARCH=amd64 go build -o app .",
            ],
        }

    def _detect_dotnet_commands(self, commands: Dict):
        """Detect .NET project run commands."""
        csproj_files = list(self.project_path.glob("**/*.csproj"))
        sln_files = list(self.project_path.glob("*.sln"))

        if not (csproj_files or sln_files):
            return

        commands["prerequisites"].extend([
            ".NET 6+ SDK installed",
        ])

        project_file = (
            str(sln_files[0].name)
            if sln_files
            else str(csproj_files[0].relative_to(self.project_path))
        )

        commands["environment_setup"] = {
            "Linux/macOS": [
                "# Install .NET SDK",
                "# https://dotnet.microsoft.com/download",
                "",
                "# Ubuntu/Debian",
                "wget https://packages.microsoft.com/config/ubuntu/"
                "22.04/packages-microsoft-prod.deb -O packages-microsoft-prod.deb",
                "sudo dpkg -i packages-microsoft-prod.deb",
                "sudo apt-get update",
                "sudo apt-get install -y dotnet-sdk-8.0",
                "",
                "# macOS",
                "brew install --cask dotnet-sdk",
                "",
                "dotnet --version",
            ],
            "Windows (PowerShell)": [
                "# Install via winget",
                "winget install Microsoft.DotNet.SDK.8",
                "",
                "# Or download from:",
                "# https://dotnet.microsoft.com/download",
                "",
                "dotnet --version",
            ],
        }
        commands["install_dependencies"] = {
            "Linux/macOS/Windows": [
                f"dotnet restore {project_file}",
            ],
        }
        commands["build_commands"] = {
            "Linux/macOS/Windows": [
                f"# Debug build",
                f"dotnet build {project_file}",
                "",
                f"# Release build",
                f"dotnet build {project_file} --configuration Release",
                "",
                f"# Publish",
                f"dotnet publish {project_file} -c Release -o ./publish",
            ],
        }
        commands["run_commands"] = {
            "Linux/macOS/Windows": [
                f"# Run in development",
                f"dotnet run --project {project_file}",
                "",
                f"# Run with specific environment",
                f"ASPNETCORE_ENVIRONMENT=Development dotnet run "
                f"--project {project_file}",
            ],
        }
        commands["test_commands"] = {
            "Linux/macOS/Windows": [
                f"dotnet test {project_file}",
                "",
                "# With coverage",
                f"dotnet test {project_file} "
                f"--collect:\"XPlat Code Coverage\"",
            ],
        }
        commands["ports"].append(
            {"port": "5000/5001", "service": "ASP.NET Core HTTP/HTTPS"}
        )

    def _detect_ruby_commands(self, commands: Dict):
        """Detect Ruby/Rails project commands."""
        gemfile = self.project_path / "Gemfile"
        if not gemfile.exists():
            return

        commands["prerequisites"].extend([
            "Ruby 3.0+ installed",
            "Bundler gem installed",
        ])

        is_rails = (
            self.project_path / "config" / "routes.rb"
        ).exists()

        commands["environment_setup"] = {
            "Linux/macOS": [
                "# Install rbenv (Ruby version manager)",
                "curl -fsSL https://github.com/rbenv/rbenv-installer/"
                "raw/HEAD/bin/rbenv-installer | bash",
                "",
                "rbenv install 3.2.0",
                "rbenv global 3.2.0",
                "",
                "# Install Bundler",
                "gem install bundler",
                "",
                "# Install project dependencies",
                "bundle install",
            ],
            "Windows (PowerShell)": [
                "# Install RubyInstaller from https://rubyinstaller.org/",
                "",
                "# Install Bundler",
                "gem install bundler",
                "",
                "# Install dependencies",
                "bundle install",
            ],
        }
        commands["install_dependencies"] = {
            "Linux/macOS/Windows": ["bundle install"],
        }

        if is_rails:
            commands["run_commands"] = {
                "Linux/macOS": [
                    "# Setup database",
                    "rails db:create",
                    "rails db:migrate",
                    "rails db:seed",
                    "",
                    "# Start Rails server",
                    "rails server",
                    "",
                    "# Or shorthand",
                    "rails s",
                    "",
                    "# Start on specific port",
                    "rails s -p 3001",
                ],
                "Windows": [
                    "rails db:create",
                    "rails db:migrate",
                    "rails server",
                ],
            }
            commands["test_commands"] = {
                "Linux/macOS/Windows": [
                    "rails test",
                    "",
                    "# RSpec (if configured)",
                    "bundle exec rspec",
                ],
            }
            commands["ports"].append(
                {"port": "3000", "service": "Rails Server"}
            )
        else:
            commands["run_commands"] = {
                "Linux/macOS/Windows": ["bundle exec ruby app.rb"],
            }

    def _detect_php_commands(self, commands: Dict):
        """Detect PHP project commands."""
        composer_json = self.project_path / "composer.json"
        if not composer_json.exists():
            return

        commands["prerequisites"].extend([
            "PHP 8.0+ installed",
            "Composer package manager",
        ])

        try:
            with open(composer_json, "r", encoding="utf-8") as f:
                composer = json.load(f)
            deps = composer.get("require", {})
            is_laravel = "laravel/framework" in deps
        except Exception:
            is_laravel = False

        commands["install_dependencies"] = {
            "Linux/macOS/Windows": ["composer install"],
        }

        if is_laravel:
            commands["environment_setup"] = {
                "Linux/macOS": [
                    "# Copy env file",
                    "cp .env.example .env",
                    "",
                    "# Install dependencies",
                    "composer install",
                    "",
                    "# Generate app key",
                    "php artisan key:generate",
                    "",
                    "# Run migrations",
                    "php artisan migrate",
                    "",
                    "# Seed database (optional)",
                    "php artisan db:seed",
                ],
                "Windows (PowerShell)": [
                    "copy .env.example .env",
                    "composer install",
                    "php artisan key:generate",
                    "php artisan migrate",
                ],
            }
            commands["run_commands"] = {
                "Linux/macOS/Windows": [
                    "php artisan serve",
                    "",
                    "# With specific port",
                    "php artisan serve --port=8080",
                ],
            }
            commands["ports"].append(
                {"port": "8000", "service": "Laravel Development Server"}
            )
        else:
            commands["run_commands"] = {
                "Linux/macOS/Windows": [
                    "php -S localhost:8080 -t public",
                ],
            }

    def _detect_rust_commands(self, commands: Dict):
        """Detect Rust project commands."""
        cargo_toml = self.project_path / "Cargo.toml"
        if not cargo_toml.exists():
            return

        commands["prerequisites"].extend([
            "Rust toolchain installed (rustup)",
        ])

        commands["environment_setup"] = {
            "Linux/macOS": [
                "# Install Rust via rustup",
                "curl --proto '=https' --tlsv1.2 -sSf "
                "https://sh.rustup.rs | sh",
                "",
                "source $HOME/.cargo/env",
                "rustup update stable",
                "",
                "rustc --version",
                "cargo --version",
            ],
            "Windows (PowerShell)": [
                "# Download rustup-init.exe from:",
                "# https://rustup.rs/",
                "",
                "rustup update stable",
                "rustc --version",
            ],
        }
        commands["install_dependencies"] = {
            "Linux/macOS/Windows": ["cargo fetch"],
        }
        commands["build_commands"] = {
            "Linux/macOS/Windows": [
                "# Debug build",
                "cargo build",
                "",
                "# Release build (optimized)",
                "cargo build --release",
            ],
        }
        commands["run_commands"] = {
            "Linux/macOS/Windows": [
                "# Run in debug mode",
                "cargo run",
                "",
                "# Run in release mode",
                "cargo run --release",
                "",
                "# Run with arguments",
                "cargo run -- --arg1 value1",
            ],
        }
        commands["test_commands"] = {
            "Linux/macOS/Windows": [
                "cargo test",
                "",
                "# Run specific test",
                "cargo test test_name",
                "",
                "# With output",
                "cargo test -- --nocapture",
            ],
        }

    def _detect_docker_commands(self, commands: Dict):
        """Detect Docker run commands."""
        dockerfile = self.project_path / "Dockerfile"
        compose_yml = self.project_path / "docker-compose.yml"
        compose_yaml = self.project_path / "docker-compose.yaml"

        has_dockerfile = dockerfile.exists()
        has_compose = compose_yml.exists() or compose_yaml.exists()
        compose_file = (
            "docker-compose.yml"
            if compose_yml.exists()
            else "docker-compose.yaml"
            if compose_yaml.exists()
            else None
        )

        if not (has_dockerfile or has_compose):
            return

        docker_cmds = {"Linux/macOS": [], "Windows (PowerShell)": []}

        if has_dockerfile:
            # Try to get project name from parent
            project_name = self.project_path.name.lower().replace(
                " ", "-"
            )
            docker_cmds["Linux/macOS"].extend([
                "# Build Docker image",
                f"docker build -t {project_name}:latest .",
                "",
                "# Build with no cache",
                f"docker build --no-cache -t {project_name}:latest .",
                "",
                "# Run Docker container",
                f"docker run -d --name {project_name} \\",
                "  -p 8080:8080 \\",
                f"  --env-file .env \\",
                f"  {project_name}:latest",
                "",
                "# View logs",
                f"docker logs -f {project_name}",
                "",
                "# Stop container",
                f"docker stop {project_name}",
                "",
                "# Remove container",
                f"docker rm {project_name}",
                "",
                "# Remove image",
                f"docker rmi {project_name}:latest",
            ])
            docker_cmds["Windows (PowerShell)"].extend([
                f"docker build -t {project_name}:latest .",
                "",
                f"docker run -d --name {project_name} `",
                "  -p 8080:8080 `",
                f"  --env-file .env `",
                f"  {project_name}:latest",
                "",
                f"docker logs -f {project_name}",
                f"docker stop {project_name}",
            ])

        if has_compose:
            compose_cmd = f"docker-compose -f {compose_file}"
            docker_cmds["Linux/macOS"].extend([
                "",
                "# ─── Docker Compose ───",
                "",
                "# Start all services",
                f"{compose_cmd} up -d",
                "",
                "# Build and start",
                f"{compose_cmd} up --build -d",
                "",
                "# View logs",
                f"{compose_cmd} logs -f",
                "",
                "# Stop all services",
                f"{compose_cmd} down",
                "",
                "# Stop and remove volumes",
                f"{compose_cmd} down -v",
                "",
                "# Scale a service",
                f"{compose_cmd} up --scale web=3 -d",
                "",
                "# Run command in service",
                f"{compose_cmd} exec web bash",
            ])
            docker_cmds["Windows (PowerShell)"].extend([
                "",
                f"{compose_cmd} up -d",
                f"{compose_cmd} up --build -d",
                f"{compose_cmd} logs -f",
                f"{compose_cmd} down",
            ])

        commands["docker_commands"] = docker_cmds

    def _detect_makefile_commands(self, commands: Dict):
        """Detect Makefile targets."""
        makefile = self.project_path / "Makefile"
        if not makefile.exists():
            return

        try:
            content = makefile.read_text(encoding="utf-8")
            # Extract targets
            targets = re.findall(
                r'^([a-zA-Z][a-zA-Z0-9_-]*)\s*:', content, re.MULTILINE
            )
            if targets:
                commands["notes"].append(
                    f"Makefile targets available: "
                    f"{', '.join(targets[:20])}"
                )
                # Common targets
                common = {
                    "install": "Install dependencies",
                    "build": "Build project",
                    "run": "Run project",
                    "start": "Start project",
                    "test": "Run tests",
                    "clean": "Clean build artifacts",
                    "docker-build": "Build Docker image",
                    "docker-run": "Run Docker container",
                    "deploy": "Deploy project",
                    "lint": "Run linter",
                    "format": "Format code",
                    "help": "Show available commands",
                }
                makefile_cmds = {
                    "Linux/macOS": [
                        "# View all available make targets",
                        "make help",
                        "",
                    ]
                }
                for target in targets:
                    if target in common:
                        makefile_cmds["Linux/macOS"].append(
                            f"make {target}  # {common[target]}"
                        )
                commands["notes"].insert(
                    0, "Makefile detected - use `make help` for all targets"
                )
        except Exception:
            pass

    def _detect_env_variables(self, commands: Dict):
        """Detect required environment variables."""
        env_files = [
            ".env.example", ".env.sample", ".env.template",
            ".env.local.example",
        ]
        for ef in env_files:
            env_file = self.project_path / ef
            if env_file.exists():
                try:
                    content = env_file.read_text(encoding="utf-8")
                    for line in content.split("\n"):
                        line = line.strip()
                        if line and not line.startswith("#"):
                            if "=" in line:
                                key = line.split("=")[0].strip()
                                if key:
                                    commands["environment_variables"].append(
                                        key
                                    )
                    if commands["environment_variables"]:
                        commands["notes"].append(
                            f"Copy {ef} to .env and fill in values"
                        )
                except Exception:
                    pass
                break

    def _detect_ports(self, commands: Dict):
        """Detect additional port configurations."""
        # Check docker-compose for ports
        compose_files = [
            self.project_path / "docker-compose.yml",
            self.project_path / "docker-compose.yaml",
        ]
        for cf in compose_files:
            if cf.exists():
                try:
                    with open(cf, "r", encoding="utf-8") as f:
                        compose = yaml.safe_load(f)
                    if isinstance(compose, dict):
                        services = compose.get("services", {})
                        for svc_name, svc_config in services.items():
                            if isinstance(svc_config, dict):
                                ports = svc_config.get("ports", [])
                                for port in ports:
                                    port_str = str(port)
                                    if ":" in port_str:
                                        host_port = port_str.split(":")[0]
                                        commands["ports"].append({
                                            "port": host_port,
                                            "service": svc_name,
                                        })
                except Exception:
                    pass

    # ----------------------------------------------------------------
    # DEVOPS PUSH & DEPLOY DETECTION
    # ----------------------------------------------------------------

    def detect_devops_push_instructions(
        self, pipeline_info: List[Dict], git_info: Dict
    ) -> Dict[str, Any]:
        """
        Generate instructions for pushing to DevOps environments.
        Covers GitHub, Azure DevOps, Jenkins, GitLab.
        """
        instructions = {
            "git_workflow": [],
            "platforms": {},
            "branch_strategy": {},
            "pre_push_checklist": [],
            "environment_promotion": [],
        }

        # Detect CI/CD platforms from pipelines
        platforms_detected = set()
        for p in pipeline_info:
            ptype = p.get("type", "")
            if ptype:
                platforms_detected.add(ptype)

        # Also detect from git remote
        remote_url = git_info.get("remote_url", "")
        if "github.com" in remote_url:
            platforms_detected.add("GitHub Actions")
        if "dev.azure.com" in remote_url or \
                "visualstudio.com" in remote_url:
            platforms_detected.add("Azure DevOps")

        # Git workflow
        instructions["git_workflow"] = self._build_git_workflow(
            git_info
        )

        # Branch strategy
        instructions["branch_strategy"] = self._build_branch_strategy(
            pipeline_info
        )

        # Platform-specific instructions
        for platform_name in platforms_detected:
            if "GitHub" in platform_name:
                instructions["platforms"]["GitHub Actions"] = \
                    self._github_push_instructions(
                        git_info, pipeline_info
                    )
            elif "Azure" in platform_name:
                instructions["platforms"]["Azure DevOps"] = \
                    self._azure_devops_push_instructions(
                        git_info, pipeline_info
                    )
            elif "Jenkins" in platform_name:
                instructions["platforms"]["Jenkins"] = \
                    self._jenkins_push_instructions(
                        git_info, pipeline_info
                    )
            elif "GitLab" in platform_name:
                instructions["platforms"]["GitLab CI"] = \
                    self._gitlab_push_instructions(
                        git_info, pipeline_info
                    )

        # If no platforms detected, provide generic instructions
        if not instructions["platforms"]:
            instructions["platforms"]["Generic Git"] = \
                self._generic_git_push_instructions(git_info)

        # Pre-push checklist
        instructions["pre_push_checklist"] = [
            "Ensure all tests pass locally before pushing",
            "Review code changes with git diff",
            "Ensure no sensitive data or secrets are committed",
            "Update documentation if APIs or configurations changed",
            "Verify environment variables are not hardcoded",
            "Check that .gitignore is properly configured",
            "Ensure commit messages follow conventions",
            "Pull latest changes before pushing to avoid conflicts",
        ]

        # Environment promotion flow
        instructions["environment_promotion"] = [
            {
                "step": "1. Development",
                "action": "Push feature branch → triggers CI build & tests",
                "branch": "feature/*, develop",
            },
            {
                "step": "2. Staging",
                "action": "Merge to staging/develop → triggers staging deployment",
                "branch": "staging, develop",
            },
            {
                "step": "3. Production",
                "action": "Merge to main/master → triggers production deployment",
                "branch": "main, master",
            },
        ]

        return instructions

    def _build_git_workflow(self, git_info: Dict) -> List[Dict]:
        """Build standard git workflow instructions."""
        current_branch = git_info.get("current_branch", "main")
        remote = "origin"

        return [
            {
                "step": "1. Check Current Status",
                "Linux/macOS/Windows": [
                    "# Check current branch and status",
                    "git status",
                    "",
                    "# Check current branch",
                    "git branch --show-current",
                    "",
                    "# View recent commits",
                    "git log --oneline -10",
                ],
            },
            {
                "step": "2. Pull Latest Changes",
                "Linux/macOS/Windows": [
                    "# Fetch all remote changes",
                    f"git fetch {remote} --prune",
                    "",
                    "# Pull latest from current branch",
                    f"git pull {remote} {current_branch}",
                    "",
                    "# Or pull with rebase (cleaner history)",
                    f"git pull --rebase {remote} {current_branch}",
                ],
            },
            {
                "step": "3. Create Feature Branch",
                "Linux/macOS/Windows": [
                    "# Create and switch to new feature branch",
                    "git checkout -b feature/your-feature-name",
                    "",
                    "# Or from a specific branch",
                    "git checkout -b feature/your-feature-name "
                    f"{current_branch}",
                ],
            },
            {
                "step": "4. Make Changes & Stage Files",
                "Linux/macOS/Windows": [
                    "# Check what changed",
                    "git diff",
                    "",
                    "# Stage specific files",
                    "git add path/to/file.ext",
                    "",
                    "# Stage all changes",
                    "git add .",
                    "",
                    "# Stage interactively (review each change)",
                    "git add -p",
                ],
            },
            {
                "step": "5. Commit Changes",
                "Linux/macOS/Windows": [
                    "# Commit with message",
                    'git commit -m "feat: add new feature description"',
                    "",
                    "# Commit types:",
                    "#   feat:     New feature",
                    "#   fix:      Bug fix",
                    "#   docs:     Documentation changes",
                    "#   style:    Formatting changes",
                    "#   refactor: Code refactoring",
                    "#   test:     Adding/fixing tests",
                    "#   chore:    Maintenance tasks",
                    "#   ci:       CI/CD changes",
                ],
            },
            {
                "step": "6. Push to Remote",
                "Linux/macOS/Windows": [
                    "# Push new branch to remote",
                    f"git push {remote} feature/your-feature-name",
                    "",
                    "# Push with upstream tracking",
                    f"git push -u {remote} feature/your-feature-name",
                    "",
                    "# Push existing branch",
                    f"git push {remote} HEAD",
                    "",
                    "# Force push (use with caution!)",
                    f"git push --force-with-lease {remote} "
                    f"feature/your-feature-name",
                ],
            },
        ]

    def _build_branch_strategy(
        self, pipeline_info: List[Dict]
    ) -> Dict:
        """Build branch strategy based on pipeline triggers."""
        triggers = []
        for p in pipeline_info:
            triggers.extend(p.get("triggers", []))

        strategy = {
            "main_branches": ["main", "master"],
            "development_branch": "develop",
            "feature_pattern": "feature/*",
            "hotfix_pattern": "hotfix/*",
            "release_pattern": "release/*",
            "description": (
                "GitFlow-based branching strategy recommended "
                "for this project"
            ),
        }

        # Detect branch names from triggers
        for trigger in triggers:
            if isinstance(trigger, str):
                if "develop" in trigger.lower():
                    strategy["development_branch"] = "develop"
                elif "staging" in trigger.lower():
                    strategy["staging_branch"] = "staging"

        return strategy

    def _github_push_instructions(
        self, git_info: Dict, pipeline_info: List[Dict]
    ) -> Dict:
        """Generate GitHub Actions push instructions."""
        remote_url = git_info.get("remote_url", "")
        # Extract owner/repo
        match = re.search(
            r'github\.com[:/]([^/]+)/([^/.]+)',
            remote_url
        )
        repo_ref = (
            f"{match.group(1)}/{match.group(2)}"
            if match else "owner/repo"
        )

        # Get workflow files
        workflow_files = [
            p["file"] for p in pipeline_info
            if "GitHub" in p.get("type", "")
        ]

        # Get triggers
        all_triggers = []
        for p in pipeline_info:
            if "GitHub" in p.get("type", ""):
                all_triggers.extend(p.get("triggers", []))

        return {
            "platform": "GitHub Actions",
            "prerequisites": [
                "GitHub account with repository access",
                "Git configured with GitHub credentials or SSH key",
                "GitHub Actions enabled on repository",
            ],
            "initial_setup": {
                "Linux/macOS": [
                    "# Configure git with GitHub credentials",
                    'git config --global user.name "Your Name"',
                    'git config --global user.email "you@example.com"',
                    "",
                    "# Option 1: Setup SSH key",
                    "ssh-keygen -t ed25519 -C 'you@example.com'",
                    "cat ~/.ssh/id_ed25519.pub",
                    "# Add the above key to GitHub:",
                    "# Settings → SSH and GPG keys → New SSH key",
                    "",
                    "# Option 2: Use GitHub CLI",
                    "gh auth login",
                    "",
                    "# Clone repository (if not already)",
                    f"git clone git@github.com:{repo_ref}.git",
                    "# OR via HTTPS:",
                    f"git clone https://github.com/{repo_ref}.git",
                ],
                "Windows (PowerShell)": [
                    'git config --global user.name "Your Name"',
                    'git config --global user.email "you@example.com"',
                    "",
                    "# Install GitHub CLI",
                    "winget install GitHub.cli",
                    "gh auth login",
                    "",
                    f"git clone https://github.com/{repo_ref}.git",
                ],
            },
            "push_workflow": {
                "Linux/macOS/Windows": [
                    "# 1. Ensure you are on feature branch",
                    "git checkout -b feature/my-feature",
                    "",
                    "# 2. Make changes, then stage and commit",
                    "git add .",
                    'git commit -m "feat: describe your changes"',
                    "",
                    "# 3. Push to GitHub",
                    "git push origin feature/my-feature",
                    "",
                    "# 4. GitHub Actions triggers automatically!",
                    f"#    Monitor at: "
                    f"https://github.com/{repo_ref}/actions",
                    "",
                    "# 5. Create Pull Request via GitHub CLI",
                    "gh pr create --title 'My Feature' \\",
                    "             --body 'Description of changes' \\",
                    "             --base main",
                    "",
                    "# Or via GitHub web UI:",
                    f"# https://github.com/{repo_ref}/compare",
                    "",
                    "# 6. After approval, merge PR",
                    "gh pr merge --merge",
                    "",
                    "# 7. This triggers the production pipeline!",
                ],
            },
            "workflow_files": workflow_files,
            "auto_triggers": all_triggers,
            "manage_secrets": {
                "Linux/macOS/Windows": [
                    "# Add secrets via GitHub CLI",
                    "gh secret set MY_SECRET --body 'secret_value'",
                    "",
                    "# List secrets",
                    "gh secret list",
                    "",
                    "# Or via GitHub Web UI:",
                    f"# https://github.com/{repo_ref}/settings/secrets/actions",
                ],
            },
            "monitor": {
                "CLI": [
                    "# Watch workflow runs",
                    "gh run list",
                    "",
                    "# Watch specific run",
                    "gh run watch",
                    "",
                    "# View run logs",
                    "gh run view --log",
                    "",
                    "# Re-run failed jobs",
                    "gh run rerun --failed",
                ],
                "Web UI": [
                    f"Pipeline Dashboard: "
                    f"https://github.com/{repo_ref}/actions",
                    f"Pull Requests: "
                    f"https://github.com/{repo_ref}/pulls",
                    f"Deployments: "
                    f"https://github.com/{repo_ref}/deployments",
                ],
            },
        }

    def _azure_devops_push_instructions(
        self, git_info: Dict, pipeline_info: List[Dict]
    ) -> Dict:
        """Generate Azure DevOps push instructions."""
        remote_url = git_info.get("remote_url", "")

        # Parse Azure DevOps URL
        org_name = "your-org"
        project_name = "your-project"
        repo_name = "your-repo"

        match = re.search(
            r'dev\.azure\.com/([^/]+)/([^/]+)/_git/([^/]+)',
            remote_url
        )
        if match:
            org_name = match.group(1)
            project_name = match.group(2)
            repo_name = match.group(3)
        else:
            match = re.search(
                r'([^.]+)\.visualstudio\.com/([^/]+)/_git/([^/]+)',
                remote_url
            )
            if match:
                org_name = match.group(1)
                project_name = match.group(2)
                repo_name = match.group(3)

        pipeline_files = [
            p["file"] for p in pipeline_info
            if "Azure" in p.get("type", "")
        ]

        return {
            "platform": "Azure DevOps",
            "prerequisites": [
                "Azure DevOps account with project access",
                "Azure CLI installed (az cli)",
                "Azure DevOps extension for Azure CLI",
                "Git configured with Azure DevOps credentials",
            ],
            "initial_setup": {
                "Linux/macOS": [
                    "# Install Azure CLI",
                    "curl -sL https://aka.ms/InstallAzureCLIDeb | sudo bash",
                    "",
                    "# macOS",
                    "brew install azure-cli",
                    "",
                    "# Install Azure DevOps extension",
                    "az extension add --name azure-devops",
                    "",
                    "# Login to Azure",
                    "az login",
                    "",
                    "# Configure default organization",
                    f"az devops configure --defaults "
                    f"organization=https://dev.azure.com/{org_name}",
                    "",
                    f"az devops configure --defaults "
                    f"project={project_name}",
                    "",
                    "# Configure git credentials",
                    "git config --global credential.helper store",
                    "",
                    "# Clone repository",
                    f"git clone https://dev.azure.com/{org_name}/"
                    f"{project_name}/_git/{repo_name}",
                ],
                "Windows (PowerShell)": [
                    "# Install Azure CLI",
                    "winget install Microsoft.AzureCLI",
                    "",
                    "# Install Azure DevOps extension",
                    "az extension add --name azure-devops",
                    "",
                    "# Login",
                    "az login",
                    "",
                    f"az devops configure --defaults "
                    f"organization=https://dev.azure.com/{org_name} "
                    f"project={project_name}",
                    "",
                    f"git clone https://dev.azure.com/{org_name}/"
                    f"{project_name}/_git/{repo_name}",
                ],
            },
            "push_workflow": {
                "Linux/macOS/Windows": [
                    "# 1. Create feature branch",
                    "git checkout -b feature/my-feature",
                    "",
                    "# 2. Make changes, commit",
                    "git add .",
                    'git commit -m "feat: describe your changes"',
                    "",
                    "# 3. Push to Azure DevOps",
                    "git push origin feature/my-feature",
                    "",
                    "# 4. Pipeline triggers automatically!",
                    f"#    Monitor at: https://dev.azure.com/"
                    f"{org_name}/{project_name}/_build",
                    "",
                    "# 5. Create Pull Request via Azure CLI",
                    "az repos pr create \\",
                    "  --title 'My Feature' \\",
                    "  --description 'Description of changes' \\",
                    "  --source-branch feature/my-feature \\",
                    "  --target-branch main \\",
                    "  --detect",
                    "",
                    "# 6. List your PRs",
                    "az repos pr list --status active",
                    "",
                    "# 7. Approve and complete PR",
                    "az repos pr update --id <PR_ID> \\",
                    "  --status completed \\",
                    "  --merge-strategy merge",
                ],
            },
            "pipeline_files": pipeline_files,
            "trigger_pipeline_manually": {
                "Linux/macOS/Windows": [
                    "# List all pipelines",
                    f"az pipelines list --project {project_name}",
                    "",
                    "# Trigger a pipeline manually",
                    f"az pipelines run --name 'your-pipeline-name' \\",
                    f"  --project {project_name}",
                    "",
                    "# Trigger with parameters",
                    f"az pipelines run --name 'your-pipeline-name' \\",
                    f"  --project {project_name} \\",
                    f"  --parameters environment=staging",
                    "",
                    "# List recent pipeline runs",
                    f"az pipelines runs list --project {project_name}",
                ],
            },
            "manage_variables": {
                "Linux/macOS/Windows": [
                    "# Create variable group",
                    "az pipelines variable-group create \\",
                    "  --name 'MyVariableGroup' \\",
                    "  --variables key1=value1 key2=value2 \\",
                    f"  --project {project_name}",
                    "",
                    "# Add secret variable",
                    "az pipelines variable-group variable create \\",
                    "  --group-id <GROUP_ID> \\",
                    "  --name MY_SECRET \\",
                    "  --value 'secret_value' \\",
                    "  --secret true",
                    "",
                    "# List variable groups",
                    f"az pipelines variable-group list "
                    f"--project {project_name}",
                ],
            },
            "monitor": {
                "CLI": [
                    "# List recent builds",
                    f"az pipelines runs list --project {project_name}",
                    "",
                    "# Get build details",
                    "az pipelines runs show --id <RUN_ID>",
                    "",
                    "# Download build artifact",
                    "az pipelines runs artifact download \\",
                    "  --run-id <RUN_ID> \\",
                    "  --artifact-name drop \\",
                    "  --path ./artifacts",
                ],
                "Web UI": [
                    f"Pipelines: https://dev.azure.com/"
                    f"{org_name}/{project_name}/_build",
                    f"Repos: https://dev.azure.com/"
                    f"{org_name}/{project_name}/_git/{repo_name}",
                    f"Releases: https://dev.azure.com/"
                    f"{org_name}/{project_name}/_release",
                    f"Boards: https://dev.azure.com/"
                    f"{org_name}/{project_name}/_boards",
                ],
            },
        }

    def _jenkins_push_instructions(
        self, git_info: Dict, pipeline_info: List[Dict]
    ) -> Dict:
        """Generate Jenkins pipeline push instructions."""
        jenkins_files = [
            p["file"] for p in pipeline_info
            if "Jenkins" in p.get("type", "")
        ]

        return {
            "platform": "Jenkins",
            "prerequisites": [
                "Jenkins server accessible",
                "Jenkins CLI or web access",
                "Git plugin installed in Jenkins",
                "Credentials configured in Jenkins",
            ],
            "initial_setup": {
                "Linux/macOS": [
                    "# Download Jenkins CLI",
                    "wget http://your-jenkins-server:8080/jnlpJars/"
                    "jenkins-cli.jar",
                    "",
                    "# Authenticate with Jenkins",
                    "java -jar jenkins-cli.jar -s "
                    "http://your-jenkins-server:8080/ \\",
                    "  -auth admin:your-api-token who-am-i",
                    "",
                    "# Or configure SSH for Jenkins",
                    "ssh-keygen -t rsa -b 4096",
                    "# Add public key to Jenkins user profile",
                ],
                "Windows (PowerShell)": [
                    "# Download Jenkins CLI JAR",
                    "Invoke-WebRequest -Uri http://your-jenkins-server:"
                    "8080/jnlpJars/jenkins-cli.jar -OutFile jenkins-cli.jar",
                    "",
                    "java -jar jenkins-cli.jar -s "
                    "http://your-jenkins-server:8080/ "
                    "-auth admin:your-api-token who-am-i",
                ],
            },
            "push_workflow": {
                "Linux/macOS/Windows": [
                    "# 1. Jenkins uses Webhooks - push to git triggers build",
                    "",
                    "# 2. Make changes and push",
                    "git add .",
                    'git commit -m "feat: your changes"',
                    "git push origin feature/my-feature",
                    "",
                    "# 3. Jenkins detects push via webhook/polling",
                    "#    and automatically starts the pipeline",
                    "",
                    "# 4. Trigger Jenkins build manually via CLI",
                    "java -jar jenkins-cli.jar \\",
                    "  -s http://your-jenkins-server:8080/ \\",
                    "  -auth admin:your-api-token \\",
                    "  build 'your-job-name' -s -v",
                    "",
                    "# 5. Trigger with parameters",
                    "java -jar jenkins-cli.jar \\",
                    "  -s http://your-jenkins-server:8080/ \\",
                    "  -auth admin:your-api-token \\",
                    "  build 'your-job-name' \\",
                    "  -p BRANCH=feature/my-feature \\",
                    "  -p ENVIRONMENT=staging -s -v",
                    "",
                    "# 6. Trigger via Jenkins REST API (curl)",
                    "curl -X POST \\",
                    "  'http://your-jenkins-server:8080/job/"
                    "your-job-name/build' \\",
                    "  --user admin:your-api-token \\",
                    "  --data-urlencode json='{\"parameter\":"
                    "[{\"name\":\"BRANCH\",\"value\":\"main\"}]}'",
                ],
            },
            "jenkinsfile_info": jenkins_files,
            "manage_credentials": {
                "Web UI Steps": [
                    "1. Navigate to Jenkins → Manage Jenkins",
                    "2. Click 'Credentials'",
                    "3. Click '(global)' → 'Add Credentials'",
                    "4. Choose credential type:",
                    "   - Username/Password for basic auth",
                    "   - SSH Username with private key for SSH",
                    "   - Secret text for API tokens",
                    "   - Secret file for key files",
                    "5. Set ID (used in Jenkinsfile)",
                    "6. Save",
                ],
            },
            "monitor": {
                "CLI": [
                    "# Get build console output",
                    "java -jar jenkins-cli.jar \\",
                    "  -s http://your-jenkins-server:8080/ \\",
                    "  -auth admin:token \\",
                    "  console 'your-job-name' <build-number>",
                    "",
                    "# List running builds",
                    "java -jar jenkins-cli.jar \\",
                    "  -s http://your-jenkins-server:8080/ \\",
                    "  -auth admin:token \\",
                    "  list-jobs",
                ],
                "Web UI": [
                    "Dashboard: http://your-jenkins-server:8080",
                    "Blue Ocean: http://your-jenkins-server:8080/blue",
                    "Job List: http://your-jenkins-server:8080/view/all",
                ],
            },
        }

    def _gitlab_push_instructions(
        self, git_info: Dict, pipeline_info: List[Dict]
    ) -> Dict:
        """Generate GitLab CI push instructions."""
        remote_url = git_info.get("remote_url", "")
        match = re.search(
            r'gitlab\.com[:/]([^/]+/[^/.]+)',
            remote_url
        )
        repo_ref = match.group(1) if match else "group/project"

        return {
            "platform": "GitLab CI",
            "prerequisites": [
                "GitLab account with project access",
                "GitLab CLI (glab) installed",
                "Git configured with GitLab credentials",
            ],
            "initial_setup": {
                "Linux/macOS": [
                    "# Install GitLab CLI",
                    "brew install glab  # macOS",
                    "",
                    "# Ubuntu/Debian",
                    "curl -s https://raw.githubusercontent.com/"
                    "profclems/glab/trunk/scripts/install.sh | sudo bash",
                    "",
                    "# Authenticate",
                    "glab auth login",
                    "",
                    f"git clone git@gitlab.com:{repo_ref}.git",
                ],
                "Windows (PowerShell)": [
                    "winget install glab",
                    "glab auth login",
                    f"git clone https://gitlab.com/{repo_ref}.git",
                ],
            },
            "push_workflow": {
                "Linux/macOS/Windows": [
                    "# 1. Create feature branch",
                    "git checkout -b feature/my-feature",
                    "",
                    "# 2. Make changes and commit",
                    "git add .",
                    'git commit -m "feat: describe changes"',
                    "",
                    "# 3. Push - GitLab CI triggers automatically",
                    "git push origin feature/my-feature",
                    "",
                    f"#    Monitor: "
                    f"https://gitlab.com/{repo_ref}/-/pipelines",
                    "",
                    "# 4. Create Merge Request via CLI",
                    "glab mr create \\",
                    "  --title 'My Feature' \\",
                    "  --description 'Changes description' \\",
                    "  --target-branch main \\",
                    "  --assignee @me",
                    "",
                    "# 5. Trigger pipeline manually",
                    "glab ci run",
                    "",
                    "# 6. View pipeline status",
                    "glab ci view",
                    "glab ci status",
                ],
            },
            "monitor": {
                "CLI": [
                    "glab ci list",
                    "glab ci view",
                    "glab ci trace  # stream logs",
                    "glab mr list --state opened",
                ],
                "Web UI": [
                    f"Pipelines: https://gitlab.com/{repo_ref}/-/pipelines",
                    f"Merge Requests: https://gitlab.com/{repo_ref}/-/merge_requests",
                ],
            },
        }

    def _generic_git_push_instructions(
        self, git_info: Dict
    ) -> Dict:
        """Generate generic git push instructions."""
        return {
            "platform": "Generic Git",
            "prerequisites": [
                "Git installed and configured",
                "Access to remote repository",
            ],
            "push_workflow": {
                "Linux/macOS/Windows": [
                    "git add .",
                    'git commit -m "your commit message"',
                    "git push origin main",
                ],
            },
        }


# ---------------------------------------------------------------------------
# Document Section: Local Dev & Deployment (ADD TO DocumentGenerator)
# ---------------------------------------------------------------------------

def add_local_dev_deployment_section(
    self,
    local_dev: Dict[str, Any],
    deploy_instructions: Dict[str, Any],
    section_number: int = 4,
):
    """
    Add Section 4: Local Development & Deployment Guide
    Insert this between Project Structure and CI/CD sections.
    """

    # Section colors rotation
    colors = ["default", "green", "purple", "orange", "berry", "teal"]

    self._add_heading(
        f"{section_number}. Local Development & Deployment Guide",
        level=1
    )
    self._add_paragraph(
        "This section provides step-by-step instructions for setting up "
        "the project locally, running it on your machine, and pushing "
        "changes to DevOps environments."
    )

    # ─── 4.1 Prerequisites ───────────────────────────────────────────
    self._add_heading(
        f"{section_number}.1 Prerequisites", level=2
    )
    prereqs = local_dev.get("prerequisites", [])
    if prereqs:
        self._add_bullet_list(list(set(prereqs)))
    else:
        self._add_paragraph("No specific prerequisites detected.")

    # Ports info
    ports = local_dev.get("ports", [])
    if ports:
        self._add_heading("Required Ports", level=3)
        port_rows = [
            [p.get("port", ""), p.get("service", "")]
            for p in ports
        ]
        self._add_table(
            ["Port", "Service"], port_rows, "teal"
        )

    # ─── 4.2 Environment Setup ───────────────────────────────────────
    self._add_heading(
        f"{section_number}.2 Environment Setup", level=2
    )

    # Env variables
    env_vars = local_dev.get("environment_variables", [])
    if env_vars:
        self._add_heading("Required Environment Variables", level=3)
        self._add_paragraph(
            "Copy the example environment file and configure the "
            "following variables:"
        )
        self._add_bullet_list(env_vars)

        env_setup_cmds = {
            "Linux/macOS": [
                "# Copy example env file",
                "cp .env.example .env",
                "",
                "# Edit with your values",
                "nano .env",
                "# or",
                "vim .env",
            ],
            "Windows (PowerShell)": [
                "# Copy example env file",
                "copy .env.example .env",
                "",
                "# Edit with Notepad",
                "notepad .env",
            ],
        }
        for platform_name, cmds in env_setup_cmds.items():
            self._add_heading(platform_name, level=3)
            self._add_code_block("\n".join(cmds))

    # Environment setup commands
    env_setup = local_dev.get("environment_setup", {})
    if env_setup:
        self._add_heading(
            "Runtime Environment Setup", level=3
        )
        for col_idx, (platform_name, cmds) in enumerate(
            env_setup.items()
        ):
            self._add_heading(
                f"🖥  {platform_name}", level=3
            )
            if isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))

    # ─── 4.3 Install Dependencies ────────────────────────────────────
    self._add_heading(
        f"{section_number}.3 Install Dependencies", level=2
    )
    install_cmds = local_dev.get("install_dependencies", {})
    if install_cmds:
        for platform_name, cmds in install_cmds.items():
            self._add_heading(platform_name, level=3)
            if isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))
    else:
        self._add_paragraph(
            "No specific dependency installation commands detected."
        )

    # ─── 4.4 Run Locally ─────────────────────────────────────────────
    self._add_heading(
        f"{section_number}.4 Run Project Locally", level=2
    )
    run_cmds = local_dev.get("run_commands", {})
    if run_cmds:
        for platform_name, cmds in run_cmds.items():
            self._add_heading(
                f"🖥  {platform_name}", level=3
            )
            if isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))
    else:
        self._add_paragraph(
            "No specific run commands detected. Please refer to "
            "the project README for instructions."
        )

    # ─── 4.5 Build Project ───────────────────────────────────────────
    build_cmds = local_dev.get("build_commands", {})
    if build_cmds:
        self._add_heading(
            f"{section_number}.5 Build Project", level=2
        )
        for platform_name, cmds in build_cmds.items():
            self._add_heading(platform_name, level=3)
            if isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))

    # ─── 4.6 Run Tests ───────────────────────────────────────────────
    test_cmds = local_dev.get("test_commands", {})
    test_section = (
        f"{section_number}.6"
        if build_cmds
        else f"{section_number}.5"
    )
    if test_cmds:
        self._add_heading(
            f"{test_section} Run Tests Locally", level=2
        )
        for platform_name, cmds in test_cmds.items():
            self._add_heading(platform_name, level=3)
            if isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))

    # ─── 4.7 Docker Local Run ────────────────────────────────────────
    docker_cmds = local_dev.get("docker_commands", {})
    docker_section = (
        f"{section_number}.7"
        if build_cmds and test_cmds
        else f"{section_number}.6"
        if build_cmds or test_cmds
        else f"{section_number}.5"
    )
    if docker_cmds:
        self._add_heading(
            f"{docker_section} Run with Docker", level=2
        )
        self._add_paragraph(
            "Docker allows you to run the project in an isolated "
            "container without installing dependencies directly on "
            "your machine."
        )
        for platform_name, cmds in docker_cmds.items():
            self._add_heading(platform_name, level=3)
            if isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))

    # ─── Notes ───────────────────────────────────────────────────────
    notes = local_dev.get("notes", [])
    if notes:
        self._add_heading("📝 Notes & Tips", level=3)
        self._add_bullet_list(notes)

    self.doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # PUSH TO DEVOPS SECTION
    # ═══════════════════════════════════════════════════════════════════
    push_section = section_number + 0.5
    self._add_heading(
        f"{section_number}b. Push & Deploy to DevOps Environments",
        level=1
    )
    self._add_paragraph(
        "This section covers the complete workflow for pushing code "
        "from your local machine to remote DevOps platforms and "
        "triggering CI/CD pipelines."
    )

    # ─── Pre-push Checklist ──────────────────────────────────────────
    self._add_heading("Pre-Push Checklist", level=2)
    self._add_paragraph(
        "Before pushing code, verify the following:"
    )
    checklist = deploy_instructions.get("pre_push_checklist", [])
    if checklist:
        for item in checklist:
            para = self.doc.add_paragraph()
            run = para.add_run(f"☐  {item}")
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.paragraph_format.left_indent = Inches(0.3)

    # ─── Git Workflow ─────────────────────────────────────────────────
    self._add_heading("Standard Git Workflow", level=2)
    git_workflow = deploy_instructions.get("git_workflow", [])
    for step_info in git_workflow:
        self._add_heading(
            step_info.get("step", ""), level=3
        )
        for platform_name, cmds in step_info.items():
            if platform_name != "step" and isinstance(cmds, list):
                self._add_code_block("\n".join(cmds))

    # ─── Branch Strategy ─────────────────────────────────────────────
    self._add_heading("Branch Strategy", level=2)
    branch_strategy = deploy_instructions.get("branch_strategy", {})
    if branch_strategy:
        branch_rows = [
            ["main / master", "Production-ready code",
             "Protected, requires PR"],
            ["develop", "Integration branch",
             "Merge feature branches here"],
            ["feature/*", "New features",
             "Branch from develop, PR to develop"],
            ["hotfix/*", "Critical bug fixes",
             "Branch from main, PR to main + develop"],
            ["release/*", "Release preparation",
             "Branch from develop, PR to main"],
        ]
        self._add_table(
            ["Branch", "Purpose", "Workflow"],
            branch_rows, "orange"
        )

    # ─── Platform-Specific Instructions ──────────────────────────────
    platforms = deploy_instructions.get("platforms", {})
    col_idx = 0
    for platform_name, platform_data in platforms.items():
        self._add_heading(
            f"Platform: {platform_name}", level=2
        )
        color = colors[col_idx % len(colors)]
        col_idx += 1

        # Prerequisites
        prereqs = platform_data.get("prerequisites", [])
        if prereqs:
            self._add_paragraph("Prerequisites:", bold=True)
            self._add_bullet_list(prereqs)

        # Initial Setup
        setup = platform_data.get("initial_setup", {})
        if setup:
            self._add_paragraph("Initial Setup:", bold=True)
            for os_name, cmds in setup.items():
                self._add_heading(f"  {os_name}", level=3)
                if isinstance(cmds, list):
                    self._add_code_block("\n".join(cmds))

        # Push Workflow
        push_wf = platform_data.get("push_workflow", {})
        if push_wf:
            self._add_paragraph(
                "Push & Deploy Workflow:", bold=True
            )
            for os_name, cmds in push_wf.items():
                self._add_heading(f"  {os_name}", level=3)
                if isinstance(cmds, list):
                    self._add_code_block("\n".join(cmds))

        # Trigger manually
        trigger = platform_data.get(
            "trigger_pipeline_manually", {}
        )
        if trigger:
            self._add_paragraph(
                "Trigger Pipeline Manually:", bold=True
            )
            for os_name, cmds in trigger.items():
                if isinstance(cmds, list):
                    self._add_code_block("\n".join(cmds))

        # Manage variables/secrets
        for section_key in [
            "manage_secrets", "manage_variables",
            "manage_credentials"
        ]:
            mgmt = platform_data.get(section_key, {})
            if mgmt:
                label = section_key.replace("_", " ").title()
                self._add_paragraph(f"{label}:", bold=True)
                for sub_key, cmds in mgmt.items():
                    if isinstance(cmds, list):
                        self._add_heading(f"  {sub_key}", level=3)
                        self._add_code_block("\n".join(cmds))

        # Monitor
        monitor = platform_data.get("monitor", {})
        if monitor:
            self._add_paragraph(
                "Monitor Pipeline & Deployments:", bold=True
            )
            monitor_rows = []
            for monitor_type, items in monitor.items():
                if isinstance(items, list):
                    for item in items:
                        if item.strip():
                            monitor_rows.append(
                                [monitor_type, item]
                            )
            if monitor_rows:
                self._add_table(
                    ["Method", "Command / URL"],
                    monitor_rows, color
                )

        # Pipeline/workflow files
        for file_key in [
            "workflow_files", "pipeline_files",
            "jenkinsfile_info"
        ]:
            files = platform_data.get(file_key, [])
            if files:
                self._add_paragraph(
                    "Pipeline Configuration Files:", bold=True
                )
                self._add_bullet_list(
                    [str(f) for f in files]
                )

        # Auto triggers
        auto_triggers = platform_data.get("auto_triggers", [])
        if auto_triggers:
            self._add_paragraph(
                "Auto Trigger Events:", bold=True
            )
            self._add_bullet_list(
                [str(t) for t in auto_triggers]
            )

    # ─── Environment Promotion Flow ───────────────────────────────────
    self._add_heading(
        "Environment Promotion Flow", level=2
    )
    promotion = deploy_instructions.get(
        "environment_promotion", []
    )
    if promotion:
        promo_rows = [
            [
                p.get("step", ""),
                p.get("action", ""),
                p.get("branch", ""),
            ]
            for p in promotion
        ]
        self._add_table(
            ["Stage", "Action", "Branch"],
            promo_rows, "berry"
        )

    # ─── Quick Reference Commands ─────────────────────────────────────
    self._add_heading("Quick Reference Commands", level=2)
    quick_ref = [
        ["git status", "Check working tree status"],
        ["git fetch --prune", "Fetch & prune remote branches"],
        ["git pull --rebase", "Pull with rebase"],
        ["git diff HEAD", "Show all uncommitted changes"],
        ["git log --oneline -10", "View last 10 commits"],
        ["git stash", "Temporarily save uncommitted changes"],
        ["git stash pop", "Restore stashed changes"],
        ["git reset --soft HEAD~1", "Undo last commit (keep changes)"],
        ["git cherry-pick <hash>", "Apply specific commit to branch"],
        ["git tag -a v1.0.0 -m 'Release'", "Create annotated tag"],
        ["git push origin --tags", "Push all tags to remote"],
        ["git branch -d feature/done", "Delete local branch"],
        ["git push origin --delete feature/done", "Delete remote branch"],
    ]
    self._add_table(
        ["Command", "Description"], quick_ref, "teal"
    )

    self.doc.add_page_break()

        # ============================================================
        # 5. CI/CD PIPELINE ANALYSIS
        # ============================================================
        self._add_heading("4. CI/CD Pipeline Analysis", level=1)

        if not pipeline_info:
            self._add_paragraph(
                "No CI/CD pipeline configuration files were detected "
                "in this project."
            )
        else:
            # 4.1 Pipeline Overview
            self._add_heading("4.1 Pipeline Overview", level=2)
            overview_rows = []
            for p in pipeline_info:
                overview_rows.append([
                    p.get("file", ""),
                    p.get("type", "Unknown"),
                    str(len(p.get("stages", []))),
                    str(len(p.get("jobs", []))),
                    str(len(p.get("triggers", []))),
                ])
            self._add_table(
                ["Pipeline File", "Type", "Stages", "Jobs",
                 "Triggers"],
                overview_rows, "default"
            )

            # 4.2 Detailed Pipeline Configuration
            self._add_heading(
                "4.2 Detailed Pipeline Configuration", level=2
            )

            color_rotation = [
                "default", "green", "purple", "orange",
                "berry", "teal",
            ]

            for p_idx, pipeline in enumerate(pipeline_info):
                color = color_rotation[p_idx % len(color_rotation)]

                self._add_heading(
                    f"Pipeline: {pipeline.get('file', 'Unknown')}",
                    level=3
                )

                # Pipeline metadata
                meta_items = [
                    f"Type: {pipeline.get('type', 'Unknown')}",
                ]
                if pipeline.get("workflow_name"):
                    meta_items.append(
                        f"Workflow Name: {pipeline['workflow_name']}"
                    )
                if pipeline.get("pool"):
                    meta_items.append(
                        f"Agent Pool: {pipeline['pool']}"
                    )
                if pipeline.get("agent"):
                    meta_items.append(
                        f"Agent: {pipeline['agent']}"
                    )
                self._add_bullet_list(meta_items)

                # Triggers
                if pipeline.get("triggers"):
                    self._add_paragraph("Triggers:", bold=True)
                    self._add_bullet_list(
                        [str(t) for t in pipeline["triggers"]]
                    )

                # Error
                if pipeline.get("error"):
                    self._add_paragraph(
                        f"⚠ Parse Warning: {pipeline['error']}",
                        italic=True
                    )

                # Raw content preview
                if pipeline.get("raw_content"):
                    self._add_paragraph(
                        "Pipeline Configuration Preview:", bold=True
                    )
                    self._add_code_block(
                        pipeline["raw_content"], max_lines=40
                    )

            # 4.3 Pipeline Stages & Jobs
            self._add_heading(
                "4.3 Pipeline Stages & Jobs", level=2
            )

            for p_idx, pipeline in enumerate(pipeline_info):
                color = color_rotation[p_idx % len(color_rotation)]

                if pipeline.get("stages") or pipeline.get("jobs"):
                    self._add_heading(
                        f"Stages/Jobs: {pipeline.get('file', '')}",
                        level=3
                    )

                    if pipeline.get("stages"):
                        self._add_paragraph("Stages:", bold=True)
                        self._add_numbered_list(
                            [str(s) if isinstance(s, str)
                             else s.get("name", str(s))
                             for s in pipeline["stages"]]
                        )

                    if pipeline.get("jobs"):
                        self._add_paragraph("Jobs Detail:", bold=True)
                        job_rows = []
                        for job in pipeline["jobs"]:
                            job_name = job.get(
                                "name",
                                job.get("displayName", "Unnamed")
                            )
                            runs_on = job.get(
                                "runs_on",
                                job.get("pool",
                                        job.get("image", ""))
                            )
                            steps_count = len(
                                job.get("steps", [])
                            )
                            env = job.get("environment", "")
                            needs = ", ".join(
                                job.get("needs", [])
                            ) if job.get("needs") else ""
                            condition = job.get("condition", "")

                            job_rows.append([
                                str(job_name),
                                str(runs_on),
                                str(steps_count),
                                str(env),
                                str(needs) if needs else str(condition),
                            ])

                        self._add_table(
                            ["Job Name", "Runs On", "Steps",
                             "Environment", "Dependencies/Conditions"],
                            job_rows, color
                        )

                        # Steps detail for each job
                        for job in pipeline["jobs"]:
                            steps = job.get("steps", [])
                            if steps:
                                job_name = job.get(
                                    "name",
                                    job.get("displayName", "Unnamed")
                                )
                                self._add_paragraph(
                                    f"Steps in '{job_name}':",
                                    bold=True
                                )
                                step_items = []
                                for step in steps:
                                    name = step.get("name", "")
                                    uses = step.get("uses", "")
                                    task = step.get("task", "")
                                    if uses:
                                        step_items.append(
                                            f"{name} (uses: {uses})"
                                        )
                                    elif task:
                                        step_items.append(
                                            f"{name} (task: {task})"
                                        )
                                    else:
                                        step_items.append(str(name))
                                self._add_bullet_list(step_items)

            # 4.4 Pipeline Variables & Parameters
            self._add_heading(
                "4.4 Pipeline Variables & Parameters", level=2
            )

            has_vars = False
            for pipeline in pipeline_info:
                if pipeline.get("variables") or \
                        pipeline.get("parameters"):
                    has_vars = True
                    self._add_heading(
                        f"Variables: {pipeline.get('file', '')}",
                        level=3
                    )

                    if pipeline.get("variables"):
                        self._add_paragraph(
                            "Environment Variables:", bold=True
                        )
                        var_rows = [
                            [k, str(v)[:60]]
                            for k, v in pipeline["variables"].items()
                        ]
                        self._add_table(
                            ["Variable", "Value"],
                            var_rows, "green"
                        )

                    if pipeline.get("parameters"):
                        self._add_paragraph(
                            "Parameters:", bold=True
                        )
                        param_rows = [
                            [
                                p.get("name", ""),
                                p.get("type", ""),
                                str(p.get("default", "")),
                            ]
                            for p in pipeline["parameters"]
                        ]
                        self._add_table(
                            ["Parameter", "Type", "Default"],
                            param_rows, "purple"
                        )

            if not has_vars:
                self._add_paragraph(
                    "No variables or parameters detected in "
                    "pipeline configurations."
                )

        self.doc.add_page_break()

        # ============================================================
        # 5. ENVIRONMENT CONFIGURATION
        # ============================================================
        self._add_heading("5. Environment Configuration", level=1)

        all_envs = set()
        for pipeline in pipeline_info:
            all_envs.update(
                e for e in pipeline.get("environments", []) if e
            )

        if all_envs:
            self._add_paragraph(
                "The following environments are referenced in the "
                "CI/CD pipelines:"
            )
            self._add_bullet_list(sorted(all_envs))
        else:
            self._add_paragraph(
                "No specific environment configurations were detected "
                "in the pipeline files."
            )

        # Check for env files
        env_files = []
        for key in file_tree:
            if ".env" in key.lower() or "environment" in key.lower():
                env_files.append(key)
        if env_files:
            self._add_heading(
                "Environment Configuration Files", level=2
            )
            self._add_bullet_list(env_files)

        # Services
        all_services = []
        for pipeline in pipeline_info:
            all_services.extend(pipeline.get("services", []))
        if all_services:
            self._add_heading("External Services & Resources", level=2)
            svc_rows = [
                [
                    s.get("type", s.get("name", "")),
                    s.get("name", str(s)),
                ]
                for s in all_services
            ]
            self._add_table(
                ["Service Type", "Service Name"],
                svc_rows, "teal"
            )

        self.doc.add_page_break()

        # ============================================================
        # 6. COMPLETE FILE LISTING
        # ============================================================
        self._add_heading("6. Complete File Listing", level=1)
        self._add_paragraph(
            "The following table lists all files in the project "
            "with their details:"
        )

        all_files_list = []
        self._collect_all_files(file_tree, all_files_list)

        if all_files_list:
            # Split into chunks if too many files
            chunk_size = 50
            for chunk_idx in range(
                0, len(all_files_list), chunk_size
            ):
                chunk = all_files_list[
                    chunk_idx:chunk_idx + chunk_size
                ]
                file_rows = [
                    [
                        f.get("path", f.get("name", "")),
                        f.get("extension", ""),
                        self._format_file_size(f.get("size", 0)),
                        f.get("modified", ""),
                    ]
                    for f in chunk
                ]
                color = ["default", "green", "purple", "orange",
                         "berry", "teal"][
                    (chunk_idx // chunk_size) % 6
                ]
                self._add_table(
                    ["File Path", "Type", "Size", "Last Modified"],
                    file_rows, color
                )

        self.doc.add_page_break()

        # ============================================================
        # 7. APPENDIX
        # ============================================================
        self._add_heading("7. Appendix", level=1)

        # 7.1 System Information
        self._add_heading("7.1 System Information", level=2)
        sys_info = [
            ["Python Version", platform.python_version()],
            ["Operating System", platform.platform()],
            ["Machine Architecture", platform.machine()],
            ["Hostname", platform.node()],
            [
                "Document Generator Version",
                "1.0.0",
            ],
        ]
        self._add_table(
            ["Property", "Value"], sys_info, "teal"
        )

        # 7.2 Glossary
        self._add_heading("7.2 Glossary", level=2)
        glossary = [
            ["CI/CD", "Continuous Integration / Continuous Deployment"],
            ["YAML", "YAML Ain't Markup Language - Configuration format"],
            ["Pipeline", "Automated workflow for building, testing, "
                         "and deploying code"],
            ["Stage", "A logical grouping of jobs in a pipeline"],
            ["Job", "A unit of work that runs on an agent/runner"],
            ["Step", "An individual task within a job"],
            ["Artifact", "Build output that can be published or "
                         "consumed"],
            ["Trigger", "An event that initiates a pipeline run"],
        ]
        self._add_table(
            ["Term", "Definition"], glossary, "purple"
        )

        # Footer note
        self._add_horizontal_line()
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(
            f"— End of Document —\n"
            f"Generated by Project Document Generator\n"
            f"{datetime.now().strftime('%B %d, %Y')}"
        )
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Save document
        self.doc.save(output_path)
        logger.info(f"Document saved to: {output_path}")

    def _render_tree(self, tree: Dict, lines: List[str],
                     prefix: str = "", is_root: bool = True):
        """Render file tree as text."""
        items = sorted(tree.keys())
        for i, key in enumerate(items):
            is_last = (i == len(items) - 1)
            value = tree[key]

            if is_root and i == 0:
                connector = ""
            elif is_last:
                connector = "└── "
            else:
                connector = "├── "

            name = Path(key).name if "/" in key or "\\" in key else key

            if isinstance(value, dict) and "name" not in value:
                # Directory
                lines.append(f"{prefix}{connector}📁 {name}/")
                new_prefix = prefix + (
                    "    " if is_last else "│   "
                )
                self._render_tree(
                    value, lines, new_prefix, False
                )
            else:
                # File
                if isinstance(value, dict):
                    size = self._format_file_size(
                        value.get("size", 0)
                    )
                    lines.append(
                        f"{prefix}{connector}📄 {name} ({size})"
                    )
                else:
                    lines.append(f"{prefix}{connector}📄 {name}")

    def _collect_all_files(self, tree: Dict,
                           result: List[Dict]):
        """Collect all file info from tree."""
        for key, value in tree.items():
            if isinstance(value, dict):
                if "name" in value and "size" in value:
                    # File info
                    value["path"] = key
                    result.append(value)
                else:
                    # Directory
                    self._collect_all_files(value, result)


# ---------------------------------------------------------------------------
# Main Orchestrator
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Project Document Generator - Scan and document "
                    "your project with AI-powered analysis",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s /path/to/project
  %(prog)s . --output my_project_doc.docx
  %(prog)s /path/to/project --output docs/project.docx
        """,
    )
    parser.add_argument(
        "project_path",
        nargs="?",
        default=".",
        help="Path to the project directory (default: current directory)",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Output document path (default: <project_name>_documentation.docx)",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose logging",
    )

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    project_path = Path(args.project_path).resolve()
    if not project_path.exists():
        logger.error(f"Project path does not exist: {project_path}")
        sys.exit(1)

    logger.info("=" * 60)
    logger.info("  PROJECT DOCUMENT GENERATOR")
    logger.info("=" * 60)
    logger.info(f"Project Path: {project_path}")
    logger.info(f"Platform: {platform.system()} {platform.release()}")
    logger.info("")

    # Step 1: Initialize components
    logger.info("[1/7] Initializing scanner and AI agent...")
    scanner = ProjectScanner(str(project_path))
    ai_agent = AIProjectAgent(str(project_path))

    # Step 2: Detect project name
    logger.info("[2/7] Detecting project name with AI agent...")
    project_name = ai_agent.detect_project_name()
    logger.info(f"  → Project Name: {project_name}")

    # Step 3: Scan file tree
    logger.info("[3/7] Scanning project file tree...")
    file_tree = scanner.scan_file_tree()
    file_stats = scanner.get_file_statistics(file_tree)
    logger.info(
        f"  → Found {file_stats['total_files']} files in "
        f"{file_stats['total_directories']} directories"
    )

    # Step 4: Detect technology stack
    logger.info("[4/7] Analyzing technology stack...")
    tech_stack = ai_agent.detect_tech_stack(file_tree)
    for category, items in tech_stack.items():
        logger.info(f"  → {category}: {', '.join(items)}")

    # Step 5: Find and parse pipeline files
    logger.info("[5/7] Scanning pipeline configurations...")
    pipeline_files = scanner.find_pipeline_files()
    pipeline_info = []
    for pf in pipeline_files:
        logger.info(f"  → Parsing: {pf.relative_to(project_path)}")
        info = scanner.parse_pipeline_file(pf)
        pipeline_info.append(info)
        if info.get("error"):
            logger.warning(f"    ⚠ {info['error']}")

    # Step 6: Get git information
    logger.info("[6/7] Collecting repository information...")
    git_info = scanner.get_git_info()
    if git_info["is_git_repo"]:
        logger.info(
            f"  → Branch: {git_info.get('current_branch', 'N/A')}"
        )
        logger.info(
            f"  → Commits: {git_info.get('total_commits', 0)}"
        )

    # Detect project type
    project_type = ai_agent.detect_project_type(tech_stack)
    logger.info(f"  → Project Type: {project_type}")

    # Generate summary
    project_summary = ai_agent.generate_project_summary(
        project_name, tech_stack, pipeline_info, file_stats
    )

    # Step 7: Generate document
    logger.info("[7/7] Generating document...")
    output_path = args.output
    if not output_path:
        safe_name = re.sub(r'[^\w\-.]', '_', project_name)
        output_path = f"{safe_name}_documentation.docx"

    doc_gen = DocumentGenerator()
    doc_gen.generate(
        project_name=project_name,
        project_type=project_type,
        project_summary=project_summary,
        tech_stack=tech_stack,
        pipeline_info=pipeline_info,
        file_tree=file_tree,
        file_stats=file_stats,
        git_info=git_info,
        project_path=str(project_path),
        output_path=output_path,
    )

    logger.info("")
    logger.info("=" * 60)
    logger.info("  DOCUMENT GENERATION COMPLETE")
    logger.info("=" * 60)
    logger.info(f"  Output: {os.path.abspath(output_path)}")
    logger.info(f"  Project: {project_name}")
    logger.info(f"  Type: {project_type}")
    logger.info(f"  Files Scanned: {file_stats['total_files']}")
    logger.info(f"  Pipelines Found: {len(pipeline_info)}")
    logger.info("=" * 60)


if __name__ == "__main__":
    main()